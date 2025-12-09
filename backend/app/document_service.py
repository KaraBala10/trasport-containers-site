"""
Document Generation Service

This module handles generation of PDF and Word documents (invoices, packing lists, etc.)
for LCL shipments.
"""

import base64
import io
import logging
import os
import re
from decimal import Decimal
from typing import Dict, List, Optional

import qrcode
from django.conf import settings
from django.template.loader import render_to_string
from django.utils import timezone
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from weasyprint import CSS, HTML
from weasyprint.text.fonts import FontConfiguration

from .models import FCLQuote, LCLShipment, PackagingPrice, Price, SyrianProvincePrice

logger = logging.getLogger(__name__)

# Try to import barcode library, but handle if not available
try:
    from barcode import Code128
    from barcode.writer import ImageWriter

    BARCODE_AVAILABLE = True
except ImportError:
    BARCODE_AVAILABLE = False
    logger.warning(
        "python-barcode library not available. Barcode generation will be disabled."
    )


def generate_barcode(shipment_number: str) -> Optional[str]:
    """
    Generate Code128 barcode for shipment number.

    Args:
        shipment_number: Shipment number to encode

    Returns:
        Base64 encoded barcode image string, or None if generation fails
    """
    if not BARCODE_AVAILABLE:
        logger.warning("Barcode library not available, skipping barcode generation")
        return None

    if not shipment_number:
        return None

    try:
        # Create Code128 barcode
        code128 = Code128(str(shipment_number), writer=ImageWriter())

        # Generate barcode image
        img_buffer = io.BytesIO()
        code128.write(
            img_buffer,
            {
                "module_width": 0.3,
                "module_height": 15.0,
                "quiet_zone": 2.0,
                "font_size": 10,
                "text_distance": 3.0,
                "background": "white",
                "foreground": "black",
            },
        )
        img_buffer.seek(0)

        # Convert to base64
        barcode_base64 = base64.b64encode(img_buffer.read()).decode("utf-8")
        return barcode_base64
    except Exception as e:
        logger.warning(f"Could not generate barcode: {str(e)}", exc_info=True)
        return None


def generate_tracking_barcode(tracking_url: str) -> Optional[str]:
    """
    Generate Code128 barcode for tracking URL.
    When scanned, the barcode will contain the tracking URL.

    Args:
        tracking_url: Tracking URL to encode (e.g., https://medo-freight.eu/tracking)

    Returns:
        Base64 encoded barcode image string, or None if generation fails
    """
    if not BARCODE_AVAILABLE:
        logger.warning("Barcode library not available, skipping barcode generation")
        return None

    if not tracking_url:
        return None

    try:
        # Create Code128 barcode with tracking URL
        code128 = Code128(str(tracking_url), writer=ImageWriter())

        # Generate barcode image with larger size for better visibility
        img_buffer = io.BytesIO()
        code128.write(
            img_buffer,
            {
                "module_width": 0.5,  # Increased from 0.3 for thicker bars
                "module_height": 25.0,  # Increased from 15.0 for taller bars
                "quiet_zone": 4.0,  # Increased from 2.0 for better scanning
                "font_size": 12,  # Increased from 8 for better readability
                "text_distance": 5.0,  # Increased from 3.0 for better spacing
                "background": "white",
                "foreground": "black",
            },
        )
        img_buffer.seek(0)

        # Convert to base64
        barcode_base64 = base64.b64encode(img_buffer.read()).decode("utf-8")
        return barcode_base64
    except Exception as e:
        logger.warning(f"Could not generate tracking barcode: {str(e)}", exc_info=True)
        return None


def calculate_invoice_totals(shipment: LCLShipment) -> Dict:
    """
    Calculate all pricing totals for invoice generation.
    Recalculates from parcels to ensure accuracy.

    Returns:
        Dict with all pricing information:
        {
            'base_lcl_price': float,
            'packaging_cost': float,
            'insurance_cost': float,
            'eu_shipping_cost': float,
            'syria_transport_cost': float,
            'total_price': float,
            'parcel_calculations': List[Dict],
            'packaging_calculations': List[Dict],
            'declared_shipment_value': float,
        }
    """
    try:
        parcels_data = shipment.parcels if shipment.parcels else []

        total_price_by_weight = 0
        total_price_by_cbm = 0
        total_packaging_cost = 0
        calculations = []
        packaging_calculations = []
        total_declared_value = 0

        # Calculate pricing for each parcel
        for parcel_data in parcels_data:
            product_id = parcel_data.get("productCategory")
            packaging_id = parcel_data.get("packagingType")
            weight = float(parcel_data.get("weight", 0))
            cbm = float(parcel_data.get("cbm", 0))
            repeat_count = int(parcel_data.get("repeatCount", 1))

            # Get shipment_type from parcel, with fallback to shipment-level
            parcel_shipment_type = parcel_data.get("shipmentType") or parcel_data.get(
                "shipment_type"
            )
            if not parcel_shipment_type and shipment.shipment_type:
                parcel_shipment_type = shipment.shipment_type
            # Ensure shipment_type is always set (even if None, it will be passed to template)
            if not parcel_shipment_type:
                parcel_shipment_type = None  # Explicitly set to None if not found

            # Collect declared value for insurance
            if parcel_data.get("wantsInsurance") or parcel_data.get(
                "isElectronicsShipment"
            ):
                declared_value = float(parcel_data.get("declaredShipmentValue", 0) or 0)
                total_declared_value += declared_value

            # Calculate packaging cost
            if packaging_id:
                try:
                    packaging = PackagingPrice.objects.get(id=int(packaging_id))
                    packaging_cost = float(packaging.price) * repeat_count
                    total_packaging_cost += packaging_cost

                    packaging_calculations.append(
                        {
                            "packaging_id": packaging_id,
                            "packaging_name_ar": packaging.ar_option,
                            "packaging_name_en": packaging.en_option,
                            "dimension": packaging.dimension,
                            "price_per_unit": float(packaging.price),
                            "repeat_count": repeat_count,
                            "total_cost": round(packaging_cost, 2),
                        }
                    )
                except PackagingPrice.DoesNotExist:
                    logger.warning(f"PackagingPrice with id {packaging_id} not found")
                except (ValueError, TypeError) as e:
                    logger.warning(f"Invalid packaging data: {str(e)}")

            # Calculate product pricing
            if product_id:
                try:
                    price = Price.objects.get(id=int(product_id))

                    if price.minimum_shipping_unit == "per_piece":
                        # Electronics: keep original calculation
                        parcel_weight = weight * repeat_count
                        parcel_cbm = cbm * repeat_count
                        price_by_weight = parcel_weight * float(price.price_per_kg)
                        price_by_cbm = parcel_cbm * float(price.price_per_kg)
                        total_price_by_weight += price_by_weight
                        total_price_by_cbm += price_by_cbm

                        calculations.append(
                            {
                                "product_id": product_id,
                                "product_name_ar": price.ar_item,
                                "product_name_en": price.en_item,
                                "weight": parcel_weight,
                                "cbm": parcel_cbm,
                                "price_per_kg": float(price.price_per_kg),
                                "price_by_weight": round(price_by_weight, 2),
                                "price_by_cbm": round(price_by_cbm, 2),
                                "is_electronics": True,
                                "hs_code": parcel_data.get("hs_code"),
                                "shipment_type": parcel_shipment_type,
                                "repeat_count": repeat_count,
                            }
                        )
                    else:
                        # Regular products: new calculation with volumetric weight
                        length = float(parcel_data.get("length", 0))
                        width = float(parcel_data.get("width", 0))
                        height = float(parcel_data.get("height", 0))

                        # Calculate volumetric weight: (L × W × H) / 6,000
                        volumetric_weight = (length * width * height) / 6000

                        # Chargeable weight = max(actual_weight, volumetric_weight)
                        chargeable_weight = max(weight, volumetric_weight)

                        # Account for repeat count
                        parcel_chargeable_weight = chargeable_weight * repeat_count

                        # Price = chargeable_weight × price_per_kg
                        parcel_price = parcel_chargeable_weight * float(
                            price.price_per_kg
                        )

                        total_price_by_weight += parcel_price
                        total_price_by_cbm += parcel_price

                        calculations.append(
                            {
                                "product_id": product_id,
                                "product_name_ar": price.ar_item,
                                "product_name_en": price.en_item,
                                "weight": weight,
                                "length": length,
                                "width": width,
                                "height": height,
                                "cbm": cbm,
                                "volumetric_weight": round(volumetric_weight, 3),
                                "chargeable_weight": round(chargeable_weight, 3),
                                "parcel_chargeable_weight": round(
                                    parcel_chargeable_weight, 3
                                ),
                                "price_per_kg": float(price.price_per_kg),
                                "price_by_weight": round(parcel_price, 2),
                                "price_by_cbm": round(parcel_price, 2),
                                "is_electronics": False,
                                "hs_code": parcel_data.get("hs_code"),
                                "shipment_type": parcel_shipment_type,
                                "repeat_count": repeat_count,
                            }
                        )
                except Price.DoesNotExist:
                    logger.warning(f"Price with id {product_id} not found")
                except (ValueError, TypeError) as e:
                    logger.warning(f"Invalid data for parcel: {str(e)}")

        # Calculate Base LCL Price: max(priceByWeight, priceByCBM, 75)
        base_lcl_price = max(total_price_by_weight, total_price_by_cbm, 75)

        # Calculate insurance if declared value exists
        insurance_cost = 0
        if total_declared_value > 0:
            # Insurance: (Base LCL Price + declared value) * 2%
            insurance_cost = (base_lcl_price + total_declared_value) * 0.02

        # Calculate EU Shipping (if exists)
        eu_shipping_cost = 0
        if (
            shipment.selected_eu_shipping_method
            and shipment.eu_pickup_weight
            and shipment.eu_pickup_weight > 0
        ):
            # EU shipping is already included in total_price, but we need to extract it
            # For now, we'll calculate it separately if needed
            # This will be handled in the template based on selected_eu_shipping_name
            pass

        # Calculate Syria Transport (if exists)
        syria_transport_cost = 0
        if (
            shipment.syria_province
            and shipment.syria_weight
            and shipment.syria_weight > 0
        ):
            try:
                province = SyrianProvincePrice.objects.get(
                    province_code=shipment.syria_province.upper(), is_active=True
                )
                syria_transport_cost = province.calculate_price(
                    float(shipment.syria_weight)
                )
            except SyrianProvincePrice.DoesNotExist:
                logger.warning(
                    f"SyrianProvincePrice for {shipment.syria_province} not found"
                )

        # Total price calculation
        calculation_total = base_lcl_price + total_packaging_cost
        total_price = calculation_total + insurance_cost

        return {
            "base_lcl_price": round(base_lcl_price, 2),
            "packaging_cost": round(total_packaging_cost, 2),
            "insurance_cost": round(insurance_cost, 2),
            "eu_shipping_cost": round(eu_shipping_cost, 2),
            "syria_transport_cost": round(syria_transport_cost, 2),
            "total_price": round(total_price, 2),
            "parcel_calculations": calculations,
            "packaging_calculations": packaging_calculations,
            "declared_shipment_value": round(total_declared_value, 2),
            "total_price_by_weight": round(total_price_by_weight, 2),
            "total_price_by_cbm": round(total_price_by_cbm, 2),
        }
    except Exception as e:
        logger.error(f"Error calculating invoice totals: {str(e)}", exc_info=True)
        raise


def get_company_info() -> Dict:
    """Get company information for invoice"""
    # Get site URL from settings or environment
    site_url = getattr(settings, "SITE_URL", "https://medo-freight.eu")
    if not site_url.startswith("http"):
        site_url = f"https://{site_url}"

    # Logo path - try multiple possible locations
    # Inside Docker: /app/WhatsApp Image...
    # Outside Docker: root directory
    possible_paths = [
        os.path.join(
            "/app", "WhatsApp Image 2025-11-28 at 23.01.45_ac5dc14b.png"
        ),  # Docker container path
        os.path.join(
            settings.BASE_DIR.parent,
            "WhatsApp Image 2025-11-28 at 23.01.45_ac5dc14b.png",
        ),  # Local development
        os.path.join(
            settings.BASE_DIR.parent.parent,
            "WhatsApp Image 2025-11-28 at 23.01.45_ac5dc14b.png",
        ),  # Alternative local path
    ]

    logo_base64 = None

    # Try to read and encode logo as base64 from any possible path
    for logo_path in possible_paths:
        if os.path.exists(logo_path):
            try:
                with open(logo_path, "rb") as f:
                    logo_data = f.read()
                    logo_base64 = base64.b64encode(logo_data).decode("utf-8")
                    logger.info(f"Successfully loaded logo from: {logo_path}")
                    break
            except Exception as e:
                logger.warning(f"Could not read logo file from {logo_path}: {str(e)}")
                continue

    if not logo_base64:
        logger.warning("Logo file not found in any of the expected locations")

    return {
        "name": "Medo-Freight EU",
        "tagline": "Ship · Route · Deliver",
        "address": "Wattweg 5, 4622RA Bergen op Zoom, Nederland",
        "phone": "+31 683083916",
        "email": "contact@medo-freight.eu",
        "website": "www.medo-freight.eu",
        "site_url": site_url,
        "logo_base64": logo_base64,
    }


def generate_invoice(shipment: LCLShipment, language: str = "ar") -> bytes:
    """
    Generate invoice PDF for LCL shipment.

    Args:
        shipment: LCLShipment instance
        language: 'ar' or 'en' (default: 'ar')

    Returns:
        PDF bytes
    """
    try:
        # Validate shipment status
        if shipment.status == "PENDING_PAYMENT":
            raise ValueError(
                "Invoice can only be generated after payment is confirmed."
            )

        if shipment.payment_status != "paid":
            raise ValueError("Payment must be confirmed before generating invoice.")

        # Calculate pricing
        pricing = calculate_invoice_totals(shipment)

        # Get company info
        company_info = get_company_info()

        # Get status display name
        from .email_service import get_status_display_name

        status_display = get_status_display_name(shipment.status, shipment.direction)

        # Calculate remaining amount
        remaining_amount = float(shipment.total_price) - float(
            shipment.amount_paid or 0
        )

        # Generate QR Code for tracking
        tracking_url = f"{company_info['site_url']}/tracking?shipment_id={shipment.id}"
        qr_code_base64 = None
        try:
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=4,
            )
            qr.add_data(tracking_url)
            qr.make(fit=True)

            # Create QR code image
            img = qr.make_image(fill_color="black", back_color="white")

            # Convert to base64
            img_buffer = io.BytesIO()
            img.save(img_buffer, format="PNG")
            img_buffer.seek(0)
            qr_code_base64 = base64.b64encode(img_buffer.read()).decode("utf-8")
        except Exception as e:
            logger.warning(f"Could not generate QR code: {str(e)}")

        # Get signature if exists
        signature_base64 = None
        if shipment.invoice_signature:
            try:
                signature_path = shipment.invoice_signature.path
                if os.path.exists(signature_path):
                    with open(signature_path, "rb") as f:
                        signature_data = f.read()
                        signature_base64 = base64.b64encode(signature_data).decode(
                            "utf-8"
                        )
                        logger.info(
                            f"Successfully loaded signature for shipment {shipment.id}"
                        )
            except Exception as e:
                logger.warning(f"Could not read signature file: {str(e)}")

        # Generate barcode for shipment number
        barcode_base64 = generate_barcode(shipment.shipment_number or str(shipment.id))

        # Prepare context for template
        context = {
            "shipment": shipment,
            "company": company_info,
            "pricing": pricing,
            "language": language,
            "invoice_date": shipment.paid_at or shipment.created_at,
            "invoice_number": "",
            "status_display": status_display,
            "remaining_amount": round(remaining_amount, 2),
            "tracking_url": tracking_url,
            "qr_code_base64": qr_code_base64,
            "barcode_base64": barcode_base64,
            "signature_base64": signature_base64,
        }

        # Render HTML template
        html_string = render_to_string("documents/invoice.html", context)

        # Generate PDF using WeasyPrint
        font_config = FontConfiguration()
        html = HTML(string=html_string, base_url=settings.BASE_DIR)

        # Generate PDF
        pdf_bytes = html.write_pdf(font_config=font_config)

        logger.info(f"Successfully generated invoice PDF for shipment {shipment.id}")
        return pdf_bytes

    except Exception as e:
        logger.error(f"Error generating invoice: {str(e)}", exc_info=True)
        raise


def generate_consolidated_export_invoice(
    shipment: LCLShipment, language: str = "en"
) -> bytes:
    """
    Generate Consolidated Export Invoice – Mixed Shipment (Personal & Commercial Goods)
    for LCL shipment. This is primarily for admin/export documentation.

    Args:
        shipment: LCLShipment instance
        language: kept for future use (currently template is EN)

    Returns:
        PDF bytes
    """
    try:
        # Reuse invoice pricing calculations (includes CBM per parcel)
        pricing = calculate_invoice_totals(shipment)

        # Company info (for logo and site URL)
        company_info = get_company_info()

        # Generate tracking URL
        tracking_url = f"{company_info['site_url']}/tracking?shipment_id={shipment.id}"

        # Generate barcode with tracking URL (scannable to go to tracking page)
        barcode_base64 = generate_tracking_barcode(
            f"{company_info['site_url']}/tracking"
        )

        # Aggregate totals for CBM, packages, and weight
        total_cbm = 0.0
        total_packages = 0
        total_weight = 0.0
        shipment_types = []

        for item in pricing.get("parcel_calculations", []):
            try:
                cbm_value = float(item.get("cbm", 0) or 0)
            except (TypeError, ValueError):
                cbm_value = 0.0
            total_cbm += cbm_value

            try:
                weight_value = float(item.get("weight", 0) or 0)
            except (TypeError, ValueError):
                weight_value = 0.0
            repeat_count = int(item.get("repeat_count", 1) or 1)
            total_weight += weight_value * repeat_count
            total_packages += repeat_count

            # Get shipment type from parcel
            if shipment.parcels:
                for parcel in shipment.parcels:
                    if isinstance(parcel, dict):
                        parcel_type = parcel.get("shipmentType") or parcel.get(
                            "shipment_type"
                        )
                        if parcel_type:
                            shipment_types.append(parcel_type)

        # Determine if shipment is personal, commercial, or mixed
        unique_types = list(set(shipment_types))
        is_personal_only = len(unique_types) == 1 and unique_types[0] == "personal"
        is_commercial_only = len(unique_types) == 1 and unique_types[0] == "commercial"
        is_mixed = len(unique_types) > 1

        context = {
            "shipment": shipment,
            "company": company_info,
            "pricing": pricing,
            "language": language,
            "invoice_date": shipment.paid_at or shipment.created_at,
            "invoice_number": "",
            "tracking_url": tracking_url,
            "barcode_base64": barcode_base64,
            "total_cbm": total_cbm,
            "total_packages": total_packages,
            "total_weight": total_weight,
            "is_personal_only": is_personal_only,
            "is_commercial_only": is_commercial_only,
            "is_mixed": is_mixed,
        }

        html_string = render_to_string(
            "documents/consolidated_export_invoice.html", context
        )

        font_config = FontConfiguration()
        html = HTML(string=html_string, base_url=settings.BASE_DIR)
        pdf_bytes = html.write_pdf(font_config=font_config)

        logger.info(
            f"Successfully generated consolidated export invoice PDF for shipment {shipment.id}"
        )
        return pdf_bytes

    except Exception as e:
        logger.error(
            f"Error generating consolidated export invoice: {str(e)}", exc_info=True
        )
        raise


def generate_packing_list(shipment: LCLShipment, language: str = "en") -> bytes:
    """
    Generate Packing List for LCL shipment.
    This document is for cargo identification purposes only.

    Args:
        shipment: LCLShipment instance
        language: kept for future use (currently template is EN)

    Returns:
        PDF bytes
    """
    try:
        # Reuse invoice pricing calculations (includes CBM per parcel)
        pricing = calculate_invoice_totals(shipment)

        # Company info (for logo and site URL)
        company_info = get_company_info()

        # Generate tracking URL
        tracking_url = f"{company_info['site_url']}/tracking?shipment_id={shipment.id}"

        # Generate barcode with tracking URL (scannable to go to tracking page)
        barcode_base64 = generate_tracking_barcode(
            f"{company_info['site_url']}/tracking"
        )

        # Aggregate totals for CBM, packages, and weight
        total_cbm = 0.0
        total_packages = 0
        total_weight = 0.0

        for item in pricing.get("parcel_calculations", []):
            try:
                cbm_value = float(item.get("cbm", 0) or 0)
            except (TypeError, ValueError):
                cbm_value = 0.0
            total_cbm += cbm_value

            try:
                weight_value = float(item.get("weight", 0) or 0)
            except (TypeError, ValueError):
                weight_value = 0.0
            repeat_count = int(item.get("repeat_count", 1) or 1)
            total_weight += weight_value * repeat_count
            total_packages += repeat_count

        # Get signature if exists
        signature_base64 = None
        if shipment.invoice_signature:
            try:
                signature_path = shipment.invoice_signature.path
                if os.path.exists(signature_path):
                    with open(signature_path, "rb") as f:
                        signature_data = f.read()
                        signature_base64 = base64.b64encode(signature_data).decode(
                            "utf-8"
                        )
                        logger.info(
                            f"Successfully loaded signature for packing list shipment {shipment.id}"
                        )
            except Exception as e:
                logger.warning(f"Could not read signature file: {str(e)}")

        context = {
            "shipment": shipment,
            "company": company_info,
            "pricing": pricing,
            "language": language,
            "invoice_date": shipment.paid_at or shipment.created_at,
            "invoice_number": shipment.shipment_number,
            "tracking_url": tracking_url,
            "barcode_base64": barcode_base64,
            "total_cbm": total_cbm,
            "total_packages": total_packages,
            "total_weight": total_weight,
            "signature_base64": signature_base64,
        }

        html_string = render_to_string("documents/packing_list.html", context)

        font_config = FontConfiguration()
        html = HTML(string=html_string, base_url=settings.BASE_DIR)
        pdf_bytes = html.write_pdf(font_config=font_config)

        logger.info(
            f"Successfully generated packing list PDF for shipment {shipment.id}"
        )
        return pdf_bytes

    except Exception as e:
        logger.error(f"Error generating packing list: {str(e)}", exc_info=True)
        raise


def generate_consolidated_packing_list(
    shipments: List[LCLShipment], language: str = "en"
) -> bytes:
    """
    Generate Consolidated Packing List for multiple LCL shipments.
    Combines common information and shows different shipments as rows.

    Args:
        shipments: List of LCLShipment instances
        language: kept for future use (currently template is EN)

    Returns:
        PDF bytes
    """
    try:
        if not shipments:
            raise ValueError("No shipments provided for consolidated packing list")

        # Company info (for logo and site URL)
        company_info = get_company_info()

        # Generate tracking URL for barcode
        tracking_url = f"{company_info['site_url']}/tracking"

        # Generate barcode with tracking URL
        barcode_base64 = generate_tracking_barcode(tracking_url)

        # Aggregate data from all shipments
        all_shipment_data = []
        grand_total_cbm = 0.0
        grand_total_packages = 0
        grand_total_weight = 0.0
        grand_total_value = 0.0

        for shipment in shipments:
            # Reuse invoice pricing calculations
            pricing = calculate_invoice_totals(shipment)

            # Calculate totals for this shipment
            total_cbm = 0.0
            total_packages = 0
            total_weight = 0.0
            total_value = 0.0

            for item in pricing.get("parcel_calculations", []):
                try:
                    cbm_value = float(item.get("cbm", 0) or 0)
                except (TypeError, ValueError):
                    cbm_value = 0.0
                total_cbm += cbm_value

                try:
                    weight_value = float(item.get("weight", 0) or 0)
                except (TypeError, ValueError):
                    weight_value = 0.0
                repeat_count = int(item.get("repeat_count", 1) or 1)
                total_weight += weight_value * repeat_count
                total_packages += repeat_count

                try:
                    price_value = float(item.get("price_by_weight", 0) or 0)
                except (TypeError, ValueError):
                    price_value = 0.0
                total_value += price_value

            grand_total_cbm += total_cbm
            grand_total_packages += total_packages
            grand_total_weight += total_weight
            grand_total_value += total_value

            # Collect all parcel items for this shipment
            shipment_items = []
            for item in pricing.get("parcel_calculations", []):
                shipment_items.append(
                    {
                        "shipment_number": shipment.shipment_number
                        or f"#{shipment.id}",
                        "product_name_en": item.get("product_name_en", ""),
                        "product_name_ar": item.get("product_name_ar", ""),
                        "hs_code": item.get("hs_code", ""),
                        "cbm": float(item.get("cbm", 0) or 0),
                        "repeat_count": int(item.get("repeat_count", 1) or 1),
                        "price_by_weight": float(item.get("price_by_weight", 0) or 0),
                        "weight": float(item.get("weight", 0) or 0),
                    }
                )

            all_shipment_data.append(
                {
                    "shipment": shipment,
                    "pricing": pricing,
                    "items": shipment_items,
                    "total_cbm": total_cbm,
                    "total_packages": total_packages,
                    "total_weight": total_weight,
                    "total_value": total_value,
                }
            )

        # Get signature from first shipment (if available)
        signature_base64 = None
        if shipments[0].invoice_signature:
            try:
                signature_path = shipments[0].invoice_signature.path
                if os.path.exists(signature_path):
                    with open(signature_path, "rb") as f:
                        signature_data = f.read()
                        signature_base64 = base64.b64encode(signature_data).decode(
                            "utf-8"
                        )
            except Exception as e:
                logger.warning(f"Could not read signature file: {str(e)}")

        context = {
            "shipments": shipments,
            "shipment_data": all_shipment_data,
            "company": company_info,
            "language": language,
            "invoice_date": timezone.now(),
            "grand_total_cbm": grand_total_cbm,
            "grand_total_packages": grand_total_packages,
            "grand_total_weight": grand_total_weight,
            "grand_total_value": grand_total_value,
            "signature_base64": signature_base64,
            "barcode_base64": barcode_base64,
            "tracking_url": tracking_url,
        }

        html_string = render_to_string(
            "documents/consolidated_packing_list.html", context
        )

        font_config = FontConfiguration()
        html = HTML(string=html_string, base_url=settings.BASE_DIR)
        pdf_bytes = html.write_pdf(font_config=font_config)

        logger.info(
            f"Successfully generated consolidated packing list PDF for {len(shipments)} shipments"
        )
        return pdf_bytes

    except Exception as e:
        logger.error(
            f"Error generating consolidated packing list: {str(e)}", exc_info=True
        )
        raise


def generate_consolidated_packing_list_bulk(
    shipments: List[LCLShipment],
    language: str = "en",
    packing_list_number: Optional[str] = None,
) -> bytes:
    """
    Generate Consolidated Packing List for multiple LCL shipments.
    Similar to generate_consolidated_export_invoice_bulk but for packing lists.

    Args:
        shipments: List of LCLShipment instances
        language: kept for future use (currently template is EN)
        packing_list_number: Optional packing list number

    Returns:
        PDF bytes
    """
    try:
        if not shipments:
            raise ValueError("No shipments provided for consolidated packing list")

        # Company info (for logo and site URL)
        company_info = get_company_info()

        # Generate tracking URL for barcode
        tracking_url = f"{company_info['site_url']}/tracking"

        # Generate barcode with tracking URL
        barcode_base64 = generate_tracking_barcode(tracking_url)

        # Aggregate data from all shipments
        all_shipment_data = []
        grand_total_cbm = 0.0
        grand_total_packages = 0
        grand_total_weight = 0.0
        grand_total_value = 0.0

        for shipment in shipments:
            # Reuse invoice pricing calculations
            pricing = calculate_invoice_totals(shipment)

            # Calculate totals for this shipment
            total_cbm = 0.0
            total_packages = 0
            total_weight = 0.0
            total_value = 0.0

            for item in pricing.get("parcel_calculations", []):
                try:
                    cbm_value = float(item.get("cbm", 0) or 0)
                except (TypeError, ValueError):
                    cbm_value = 0.0
                total_cbm += cbm_value

                try:
                    weight_value = float(item.get("weight", 0) or 0)
                except (TypeError, ValueError):
                    weight_value = 0.0
                repeat_count = int(item.get("repeat_count", 1) or 1)
                total_weight += weight_value * repeat_count
                total_packages += repeat_count

                try:
                    price_value = float(item.get("price_by_weight", 0) or 0)
                except (TypeError, ValueError):
                    price_value = 0.0
                total_value += price_value

            grand_total_cbm += total_cbm
            grand_total_packages += total_packages
            grand_total_weight += total_weight
            grand_total_value += total_value

            # Collect all parcel items for this shipment
            shipment_items = []
            for item in pricing.get("parcel_calculations", []):
                shipment_items.append(
                    {
                        "shipment_number": shipment.shipment_number
                        or f"#{shipment.id}",
                        "product_name_en": item.get("product_name_en", ""),
                        "product_name_ar": item.get("product_name_ar", ""),
                        "hs_code": item.get("hs_code", ""),
                        "cbm": float(item.get("cbm", 0) or 0),
                        "repeat_count": int(item.get("repeat_count", 1) or 1),
                        "price_by_weight": float(item.get("price_by_weight", 0) or 0),
                        "weight": float(item.get("weight", 0) or 0),
                    }
                )

            all_shipment_data.append(
                {
                    "shipment": shipment,
                    "pricing": pricing,
                    "items": shipment_items,
                    "total_cbm": total_cbm,
                    "total_packages": total_packages,
                    "total_weight": total_weight,
                    "total_value": total_value,
                }
            )

        # Generate packing list number if not provided
        if not packing_list_number:
            packing_list_number = (
                f"PL-{timezone.now().strftime('%Y%m%d')}-{len(shipments)}"
            )

        context = {
            "shipments": shipments,
            "shipment_data": all_shipment_data,
            "company": company_info,
            "language": language,
            "invoice_date": timezone.now(),
            "packing_list_number": packing_list_number,
            "grand_total_cbm": grand_total_cbm,
            "grand_total_packages": grand_total_packages,
            "grand_total_weight": grand_total_weight,
            "grand_total_value": grand_total_value,
            "barcode_base64": barcode_base64,
            "tracking_url": tracking_url,
        }

        html_string = render_to_string(
            "documents/consolidated_packing_list.html", context
        )

        font_config = FontConfiguration()
        html = HTML(string=html_string, base_url=settings.BASE_DIR)
        pdf_bytes = html.write_pdf(font_config=font_config)

        logger.info(
            f"Successfully generated consolidated packing list PDF for {len(shipments)} shipments"
        )
        return pdf_bytes

    except Exception as e:
        logger.error(
            f"Error generating consolidated packing list bulk: {str(e)}",
            exc_info=True,
        )
        raise


def generate_multiple_consolidated_packing_lists(
    shipments: List[LCLShipment],
    language: str = "en",
    max_cbm: float = 65.0,
    max_weight_kg: float = 24000.0,
) -> bytes:
    """
    Generate multiple consolidated packing lists split by CBM and weight limits,
    and return them as a ZIP file.

    Args:
        shipments: List of LCLShipment instances
        language: Language for packing lists (default: 'en')
        max_cbm: Maximum CBM per packing list (default: 65.0)
        max_weight_kg: Maximum weight per packing list (default: 24000.0)

    Returns:
        ZIP file bytes containing all packing lists
    """
    import zipfile
    from datetime import datetime

    try:
        # Split shipments into groups
        groups = split_shipments_by_limits(shipments, max_cbm, max_weight_kg)

        if not groups:
            raise ValueError("No shipment groups created")

        # Create ZIP file in memory
        zip_buffer = io.BytesIO()
        date_str = datetime.now().strftime("%Y%m%d")

        try:
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                for idx, group in enumerate(groups, start=1):
                    # Generate packing list number for this group
                    packing_list_number = f"PL-{date_str}-{idx:03d}"

                    # Generate PDF for this group
                    pdf_bytes = generate_consolidated_packing_list_bulk(
                        group,
                        language=language,
                        packing_list_number=packing_list_number,
                    )

                    # Add PDF to ZIP
                    filename = f"Consolidated-Packing-List-{packing_list_number}.pdf"
                    zip_file.writestr(filename, pdf_bytes)

                    logger.info(
                        f"Generated packing list {packing_list_number} with {len(group)} shipments"
                    )

            # Ensure ZIP file is properly closed and get bytes
            zip_buffer.seek(0)
            zip_bytes = zip_buffer.getvalue()

            # Verify ZIP file is valid
            if len(zip_bytes) == 0:
                raise ValueError("Generated ZIP file is empty")

            # Test ZIP file integrity
            test_zip = zipfile.ZipFile(io.BytesIO(zip_bytes), "r")
            test_zip.close()

            logger.info(
                f"ZIP file created successfully, size: {len(zip_bytes)} bytes, files: {len(groups)}"
            )

        except Exception as zip_error:
            logger.error(f"Error creating ZIP file: {str(zip_error)}", exc_info=True)
            raise

        logger.info(
            f"Successfully generated {len(groups)} packing lists and compressed into ZIP file"
        )
        return zip_bytes

    except Exception as e:
        logger.error(
            f"Error generating multiple consolidated packing lists: {str(e)}",
            exc_info=True,
        )
        raise


def generate_consolidated_export_invoice_bulk(
    shipments: List[LCLShipment],
    language: str = "en",
    invoice_number: Optional[str] = None,
) -> bytes:
    """
    Generate Consolidated Export Invoice for multiple LCL shipments.
    Combines common information and shows different shipments as rows.

    Args:
        shipments: List of LCLShipment instances
        language: kept for future use (currently template is EN)

    Returns:
        PDF bytes
    """
    try:
        if not shipments:
            raise ValueError("No shipments provided for consolidated export invoice")

        # Company info (for logo and site URL)
        company_info = get_company_info()

        # Aggregate data from all shipments
        all_shipment_data = []
        grand_total_cbm = 0.0
        grand_total_packages = 0
        grand_total_weight = 0.0
        grand_total_value = 0.0
        all_shipment_types = []

        for shipment in shipments:
            # Reuse invoice pricing calculations
            pricing = calculate_invoice_totals(shipment)

            # Calculate totals for this shipment
            total_cbm = 0.0
            total_packages = 0
            total_weight = 0.0
            total_value = 0.0
            shipment_types = []

            for item in pricing.get("parcel_calculations", []):
                try:
                    cbm_value = float(item.get("cbm", 0) or 0)
                except (TypeError, ValueError):
                    cbm_value = 0.0
                total_cbm += cbm_value

                try:
                    weight_value = float(item.get("weight", 0) or 0)
                except (TypeError, ValueError):
                    weight_value = 0.0
                repeat_count = int(item.get("repeat_count", 1) or 1)
                total_weight += weight_value * repeat_count
                total_packages += repeat_count

                try:
                    price_value = float(item.get("price_by_weight", 0) or 0)
                except (TypeError, ValueError):
                    price_value = 0.0
                total_value += price_value

            # Get shipment types from parcels
            if shipment.parcels:
                for parcel in shipment.parcels:
                    if isinstance(parcel, dict):
                        parcel_type = parcel.get("shipmentType") or parcel.get(
                            "shipment_type"
                        )
                        if parcel_type:
                            shipment_types.append(parcel_type)

            all_shipment_types.extend(shipment_types)

            grand_total_cbm += total_cbm
            grand_total_packages += total_packages
            grand_total_weight += total_weight
            grand_total_value += total_value

            # Collect all parcel items for this shipment
            shipment_items = []
            for idx, item in enumerate(pricing.get("parcel_calculations", [])):
                # Try to get shipment_type from item first, then fallback to shipment-level
                shipment_type = item.get("shipment_type") or None

                # If not found in item, try to get from shipment.parcels directly
                if not shipment_type and shipment.parcels:
                    try:
                        # Match by index if possible
                        if idx < len(shipment.parcels):
                            parcel = shipment.parcels[idx]
                            if isinstance(parcel, dict):
                                shipment_type = parcel.get(
                                    "shipmentType"
                                ) or parcel.get("shipment_type")
                    except (IndexError, TypeError):
                        pass

                # Final fallback: use shipment-level shipment_type
                if not shipment_type and shipment.shipment_type:
                    shipment_type = shipment.shipment_type

                # Log for debugging
                logger.info(
                    f"Shipment {shipment.id} item {idx}: shipment_type={shipment_type}, "
                    f"price_by_weight={item.get('price_by_weight', 0)}"
                )

                shipment_items.append(
                    {
                        "shipment_number": shipment.shipment_number
                        or f"#{shipment.id}",
                        "product_name_en": item.get("product_name_en", ""),
                        "product_name_ar": item.get("product_name_ar", ""),
                        "hs_code": item.get("hs_code", ""),
                        "cbm": float(item.get("cbm", 0) or 0),
                        "repeat_count": int(item.get("repeat_count", 1) or 1),
                        "price_by_weight": float(item.get("price_by_weight", 0) or 0),
                        "price_per_kg": float(item.get("price_per_kg", 0) or 0),
                        "weight": float(item.get("weight", 0) or 0),
                        "shipment_type": shipment_type,  # personal or commercial
                    }
                )

            # Determine shipment type
            unique_types = list(set(shipment_types))
            is_personal_only = len(unique_types) == 1 and unique_types[0] == "personal"
            is_commercial_only = (
                len(unique_types) == 1 and unique_types[0] == "commercial"
            )
            is_mixed = len(unique_types) > 1

            all_shipment_data.append(
                {
                    "shipment": shipment,
                    "pricing": pricing,
                    "items": shipment_items,
                    "total_cbm": total_cbm,
                    "total_packages": total_packages,
                    "total_weight": total_weight,
                    "total_value": total_value,
                    "is_personal_only": is_personal_only,
                    "is_commercial_only": is_commercial_only,
                    "is_mixed": is_mixed,
                }
            )

        # Determine overall type
        unique_all_types = list(set(all_shipment_types))
        overall_is_personal_only = (
            len(unique_all_types) == 1 and unique_all_types[0] == "personal"
        )
        overall_is_commercial_only = (
            len(unique_all_types) == 1 and unique_all_types[0] == "commercial"
        )
        overall_is_mixed = len(unique_all_types) > 1

        # Generate invoice number if not provided
        if not invoice_number:
            invoice_number = f"INV-{timezone.now().strftime('%Y%m%d')}-{len(shipments)}"

        # Generate tracking URL for barcode
        tracking_url = f"{company_info['site_url']}/tracking"

        # Generate barcode with tracking URL
        barcode_base64 = generate_tracking_barcode(tracking_url)

        context = {
            "shipments": shipments,
            "shipment_data": all_shipment_data,
            "company": company_info,
            "language": language,
            "invoice_date": timezone.now(),
            "invoice_number": invoice_number,
            "grand_total_cbm": grand_total_cbm,
            "grand_total_packages": grand_total_packages,
            "grand_total_weight": grand_total_weight,
            "grand_total_value": grand_total_value,
            "is_personal_only": overall_is_personal_only,
            "is_commercial_only": overall_is_commercial_only,
            "is_mixed": overall_is_mixed,
            "barcode_base64": barcode_base64,
            "tracking_url": tracking_url,
        }

        html_string = render_to_string(
            "documents/consolidated_export_invoice.html", context
        )

        font_config = FontConfiguration()
        html = HTML(string=html_string, base_url=settings.BASE_DIR)
        pdf_bytes = html.write_pdf(font_config=font_config)

        logger.info(
            f"Successfully generated consolidated export invoice PDF for {len(shipments)} shipments"
        )
        return pdf_bytes

    except Exception as e:
        logger.error(
            f"Error generating consolidated export invoice bulk: {str(e)}",
            exc_info=True,
        )
        raise


def split_shipments_by_limits(
    shipments: List[LCLShipment],
    max_cbm: float = 65.0,
    max_weight_kg: float = 24000.0,
) -> List[List[LCLShipment]]:
    """
    Split shipments into groups based on CBM and weight limits.
    Each group will not exceed max_cbm or max_weight_kg.
    If a single shipment exceeds the limits, it will be duplicated across multiple groups
    to ensure each invoice stays within limits.

    Args:
        shipments: List of LCLShipment instances
        max_cbm: Maximum CBM per group (default: 65.0)
        max_weight_kg: Maximum weight in kg per group (default: 24000.0)

    Returns:
        List of shipment groups, where each group is a list of shipments
    """
    groups = []
    current_group = []
    current_cbm = 0.0
    current_weight = 0.0

    for shipment in shipments:
        # Calculate totals for this shipment
        try:
            pricing = calculate_invoice_totals(shipment)
            shipment_cbm = 0.0
            shipment_weight = 0.0

            for item in pricing.get("parcel_calculations", []):
                try:
                    cbm_value = float(item.get("cbm", 0) or 0)
                except (TypeError, ValueError):
                    cbm_value = 0.0
                shipment_cbm += cbm_value

                try:
                    weight_value = float(item.get("weight", 0) or 0)
                except (TypeError, ValueError):
                    weight_value = 0.0
                repeat_count = int(item.get("repeat_count", 1) or 1)
                shipment_weight += weight_value * repeat_count

            # If this single shipment exceeds limits, we need to split it across multiple groups
            # Calculate how many groups this shipment needs
            if shipment_cbm > max_cbm or shipment_weight > max_weight_kg:
                # Calculate number of groups needed for this shipment
                groups_needed_cbm = int(shipment_cbm / max_cbm) + (
                    1 if shipment_cbm % max_cbm > 0 else 0
                )
                groups_needed_weight = int(shipment_weight / max_weight_kg) + (
                    1 if shipment_weight % max_weight_kg > 0 else 0
                )
                groups_needed = max(groups_needed_cbm, groups_needed_weight)

                # If current group has items, finalize it first
                if current_group:
                    groups.append(current_group)
                    current_group = []
                    current_cbm = 0.0
                    current_weight = 0.0

                # Distribute this shipment across multiple groups
                # Each group gets a copy of the shipment (they will show the same data but different invoice numbers)
                for i in range(groups_needed):
                    groups.append([shipment])

                # Reset current group since we've distributed this shipment
                current_group = []
                current_cbm = 0.0
                current_weight = 0.0
                continue

            # Check if adding this shipment would exceed limits
            if current_group and (
                current_cbm + shipment_cbm > max_cbm
                or current_weight + shipment_weight > max_weight_kg
            ):
                # Start a new group
                groups.append(current_group)
                current_group = [shipment]
                current_cbm = shipment_cbm
                current_weight = shipment_weight
            else:
                # Add to current group
                current_group.append(shipment)
                current_cbm += shipment_cbm
                current_weight += shipment_weight

        except Exception as e:
            logger.error(
                f"Error calculating totals for shipment {shipment.id}: {str(e)}"
            )
            # Add shipment anyway to avoid losing data
            if current_group:
                groups.append(current_group)
            current_group = [shipment]
            current_cbm = 0.0
            current_weight = 0.0

    # Add the last group if it has shipments
    if current_group:
        groups.append(current_group)

    logger.info(
        f"Split {len(shipments)} shipments into {len(groups)} groups based on CBM ({max_cbm}) and weight ({max_weight_kg} kg) limits"
    )
    return groups


def generate_multiple_consolidated_invoices(
    shipments: List[LCLShipment],
    language: str = "en",
    max_cbm: float = 65.0,
    max_weight_kg: float = 24000.0,
) -> bytes:
    """
    Generate multiple consolidated export invoices split by CBM and weight limits,
    and return them as a ZIP file.

    Args:
        shipments: List of LCLShipment instances
        language: Language for invoices (default: 'en')
        max_cbm: Maximum CBM per invoice (default: 65.0)
        max_weight_kg: Maximum weight per invoice (default: 24000.0)

    Returns:
        ZIP file bytes containing all invoices
    """
    import zipfile
    from datetime import datetime

    try:
        # Split shipments into groups
        groups = split_shipments_by_limits(shipments, max_cbm, max_weight_kg)

        if not groups:
            raise ValueError("No shipment groups created")

        # Create ZIP file in memory
        zip_buffer = io.BytesIO()
        date_str = datetime.now().strftime("%Y%m%d")

        try:
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                for idx, group in enumerate(groups, start=1):
                    # Generate invoice number for this group
                    invoice_number = f"INV-{date_str}-{idx:03d}"

                    # Generate PDF for this group
                    pdf_bytes = generate_consolidated_export_invoice_bulk(
                        group, language=language, invoice_number=invoice_number
                    )

                    # Add PDF to ZIP
                    filename = f"Consolidated-Export-Invoice-{invoice_number}.pdf"
                    zip_file.writestr(filename, pdf_bytes)

                    logger.info(
                        f"Generated invoice {invoice_number} with {len(group)} shipments"
                    )

            # Ensure ZIP file is properly closed and get bytes
            zip_buffer.seek(0)
            zip_bytes = zip_buffer.getvalue()

            # Verify ZIP file is valid
            if len(zip_bytes) == 0:
                raise ValueError("Generated ZIP file is empty")

            # Test ZIP file integrity
            test_zip = zipfile.ZipFile(io.BytesIO(zip_bytes), "r")
            test_zip.close()

            logger.info(
                f"ZIP file created successfully, size: {len(zip_bytes)} bytes, files: {len(groups)}"
            )

        except Exception as zip_error:
            logger.error(f"Error creating ZIP file: {str(zip_error)}", exc_info=True)
            raise

        logger.info(
            f"Successfully generated {len(groups)} invoices and compressed into ZIP file"
        )
        return zip_bytes

    except Exception as e:
        logger.error(
            f"Error generating multiple consolidated invoices: {str(e)}",
            exc_info=True,
        )
        raise


def generate_shipping_labels(
    shipment: LCLShipment, language: str = "ar", num_labels: Optional[int] = None
) -> bytes:
    """
    Generate shipping labels PDF for LCL shipment.
    Creates one label per parcel (including repeat_count).

    Args:
        shipment: LCLShipment instance
        language: 'ar' or 'en' (default: 'ar')

    Returns:
        PDF bytes containing all shipping labels
    """
    try:
        # Validate shipment
        if not shipment:
            raise ValueError("Shipment is required")

        # Get company info
        company_info = get_company_info()
        if not company_info or not isinstance(company_info, dict):
            raise ValueError("Company info is not available")

        # Generate QR Code for tracking
        site_url = company_info.get("site_url", "https://medo-freight.eu")
        tracking_url = f"{site_url}/tracking?shipment_id={shipment.id}"
        qr_code_base64 = None
        try:
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=8,
                border=2,
            )
            qr.add_data(tracking_url)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white")
            img_buffer = io.BytesIO()
            img.save(img_buffer, format="PNG")
            img_buffer.seek(0)
            qr_code_base64 = base64.b64encode(img_buffer.read()).decode("utf-8")
        except Exception as e:
            logger.warning(f"Could not generate QR code: {str(e)}")

        # Generate barcode for shipment number
        barcode_base64 = None
        try:
            shipment_number = (
                shipment.shipment_number or str(shipment.id) if shipment else "000000"
            )
            barcode_base64 = generate_barcode(shipment_number)
            if not barcode_base64:
                logger.warning(
                    f"Barcode generation returned None for shipment {shipment.id if shipment else 'unknown'}"
                )
        except Exception as barcode_error:
            logger.warning(f"Could not generate barcode: {str(barcode_error)}")
            barcode_base64 = None

        # Calculate total number of labels needed
        parcels_data = shipment.parcels if shipment.parcels else []

        # If num_labels is provided, use it; otherwise calculate from parcels
        if num_labels is not None and num_labels > 0:
            total_labels = num_labels
            # Use first parcel data for all labels
            base_parcel = {}
            if parcels_data:
                first_parcel = parcels_data[0]
                if first_parcel and isinstance(first_parcel, dict):
                    base_parcel = first_parcel.copy()

            label_parcels = []
            for i in range(total_labels):
                label_parcels.append(
                    {
                        **base_parcel,
                        "parcel_index": i + 1,
                        "total_parcels": total_labels,
                    }
                )
        else:
            # Calculate from parcels (original logic)
            total_labels = 0
            label_parcels = []  # List of parcels with their repeat counts

            for parcel_data in parcels_data:
                # Skip None or invalid parcel data
                if not parcel_data or not isinstance(parcel_data, dict):
                    logger.warning(f"Skipping invalid parcel data: {parcel_data}")
                    continue

                repeat_count = int(parcel_data.get("repeatCount", 1) or 1)
                total_labels += repeat_count
                # Add this parcel repeat_count times
                for i in range(repeat_count):
                    label_parcels.append(
                        {
                            **parcel_data,
                            "parcel_index": i + 1,
                            "total_parcels": repeat_count,
                        }
                    )

            if total_labels == 0:
                raise ValueError("No parcels found in shipment")

        # Generate HTML for all labels
        all_labels_html = []

        for idx, parcel in enumerate(label_parcels, 1):
            # Ensure parcel is always a dict
            if not parcel or not isinstance(parcel, dict):
                parcel = {}

            # Ensure all context values are safe
            context = {
                "shipment": shipment,
                "company": company_info if company_info else {},
                "parcel": parcel if parcel else {},
                "parcel_index": idx,
                "total_labels": total_labels,
                "language": language or "en",
                "qr_code_base64": qr_code_base64 or "",
                "barcode_base64": barcode_base64 or "",
            }

            # Render label template
            try:
                label_html = render_to_string("documents/shipping_label.html", context)
                if label_html and len(label_html.strip()) > 0:
                    all_labels_html.append(label_html)
                else:
                    logger.warning(f"Empty HTML generated for label {idx}, skipping")
            except Exception as render_error:
                logger.error(
                    f"Error rendering label {idx}: {str(render_error)}", exc_info=True
                )
                # Continue with next label
                continue

        # Combine all labels into one HTML document
        # Extract CSS from first label (remove @page to avoid issues)
        css_content = ""
        if all_labels_html and len(all_labels_html) > 0:
            try:
                style_match = re.search(
                    r"<style[^>]*>(.*?)</style>",
                    all_labels_html[0],
                    re.DOTALL | re.IGNORECASE,
                )
                if style_match and style_match.group(1):
                    css_content = style_match.group(1)
                    # Remove @page rules completely - we'll add it via stylesheets
                    css_content = re.sub(
                        r"@page\s*\{[^}]*\}",
                        "",
                        css_content,
                        flags=re.DOTALL | re.IGNORECASE,
                    )
            except Exception as e:
                logger.warning(f"Could not extract CSS from label HTML: {str(e)}")
                css_content = ""

        # Build combined HTML - each label on separate page
        # Build HTML string carefully to avoid CSS curly brace issues
        combined_html = "<!DOCTYPE html>\n"
        combined_html += "<html>\n"
        combined_html += "<head>\n"
        combined_html += '    <meta charset="UTF-8">\n'
        combined_html += "    <style>\n"
        combined_html += css_content + "\n"
        combined_html += "    </style>\n"
        combined_html += "</head>\n"
        combined_html += "<body>\n"

        # Add each label's body content
        for idx, label_html in enumerate(all_labels_html, 1):
            try:
                # Extract body content
                body_match = re.search(
                    r"<body[^>]*>(.*?)</body>", label_html, re.DOTALL | re.IGNORECASE
                )
                if body_match and body_match.group(1):
                    body_content = body_match.group(1).strip()
                    # Add page break except for last label
                    if idx < len(all_labels_html):
                        combined_html += f'<div style="page-break-after: always;">{body_content}</div>'
                    else:
                        combined_html += f"<div>{body_content}</div>"
                else:
                    # Fallback: extract label-container
                    container_match = re.search(
                        r'<div[^>]*class=["\']label-container["\'][^>]*>(.*?)</div>',
                        label_html,
                        re.DOTALL | re.IGNORECASE,
                    )
                    if container_match and container_match.group(1):
                        if idx < len(all_labels_html):
                            combined_html += f'<div style="page-break-after: always;"><div class="label-container">{container_match.group(1).strip()}</div></div>'
                        else:
                            combined_html += f'<div><div class="label-container">{container_match.group(1).strip()}</div></div>'
                    else:
                        # Last resort: use the entire HTML
                        if idx < len(all_labels_html):
                            combined_html += f'<div style="page-break-after: always;">{label_html}</div>'
                        else:
                            combined_html += f"<div>{label_html}</div>"
            except Exception as e:
                logger.error(f"Error processing label {idx}: {str(e)}", exc_info=True)
                # Skip this label and continue
                continue

        combined_html += "</body>\n"
        combined_html += "</html>"

        # Validate combined HTML
        if not combined_html or len(combined_html.strip()) < 100:
            raise ValueError("Generated HTML is empty or too short")

        # Validate that we have labels
        if not all_labels_html or len(all_labels_html) == 0:
            raise ValueError("No labels were generated")

        # Generate PDF using WeasyPrint with 6x4 inch page size
        # Use CSS stylesheet to set page size (avoids extra_skip_height bug)
        try:
            font_config = FontConfiguration()

            # Ensure base_url is valid
            base_url = settings.BASE_DIR if hasattr(settings, "BASE_DIR") else None
            if not base_url:
                base_url = os.path.dirname(os.path.abspath(__file__))

            # Ensure base_url is a string path
            if isinstance(base_url, str):
                base_url = base_url
            else:
                base_url = (
                    str(base_url)
                    if base_url
                    else os.path.dirname(os.path.abspath(__file__))
                )

            html = HTML(string=combined_html, base_url=base_url)

            # Create CSS for 6x4 inch page size
            page_css = None
            try:
                page_css = CSS(string="@page { size: 6in 4in; margin: 0; }")
            except Exception as css_error:
                logger.warning(f"Could not create CSS stylesheet: {str(css_error)}")
                # Continue without custom CSS - WeasyPrint will use default

            # Generate PDF with custom page size via stylesheet
            try:
                if page_css:
                    pdf_bytes = html.write_pdf(
                        font_config=font_config, stylesheets=[page_css]
                    )
                else:
                    pdf_bytes = html.write_pdf(font_config=font_config)
            except Exception as write_error:
                logger.error(
                    f"Error in html.write_pdf: {str(write_error)}", exc_info=True
                )
                # Try without stylesheet as fallback
                try:
                    pdf_bytes = html.write_pdf(font_config=font_config)
                except Exception as fallback_error:
                    logger.error(
                        f"Error in fallback html.write_pdf: {str(fallback_error)}",
                        exc_info=True,
                    )
                    raise ValueError(f"Failed to generate PDF: {str(fallback_error)}")

            if not pdf_bytes or len(pdf_bytes) == 0:
                raise ValueError("Generated PDF is empty")

            logger.info(
                f"Successfully generated {total_labels} shipping labels for shipment {shipment.id} ({len(pdf_bytes)} bytes)"
            )
            return pdf_bytes
        except ValueError:
            # Re-raise ValueError as-is
            raise
        except Exception as pdf_error:
            logger.error(f"Error generating PDF: {str(pdf_error)}", exc_info=True)
            # Log the full traceback for debugging
            import traceback

            logger.error(f"Full traceback: {traceback.format_exc()}")
            raise ValueError(f"Failed to generate PDF: {str(pdf_error)}")

    except Exception as e:
        logger.error(f"Error generating shipping labels: {str(e)}", exc_info=True)
        raise


def save_invoice_to_storage(shipment: LCLShipment, pdf_bytes: bytes) -> str:
    """
    Save invoice PDF to storage and update shipment record.

    Args:
        shipment: LCLShipment instance
        pdf_bytes: PDF file bytes

    Returns:
        File path
    """
    try:
        from django.core.files.base import ContentFile

        # Generate filename
        filename = f"Invoice-{shipment.shipment_number}.pdf"

        # Save to FileField
        shipment.invoice_file.save(filename, ContentFile(pdf_bytes), save=False)
        shipment.invoice_generated_at = timezone.now()
        shipment.save()

        logger.info(
            f"Saved invoice to {shipment.invoice_file.path} for shipment {shipment.id}"
        )
        return shipment.invoice_file.path

    except Exception as e:
        logger.error(f"Error saving invoice to storage: {str(e)}", exc_info=True)
        raise


def generate_receipt(shipment: LCLShipment, language: str = "en") -> bytes:
    """
    Generate receipt PDF for LCL shipment.

    Args:
        shipment: LCLShipment instance
        language: 'ar' or 'en' (default: 'en')

    Returns:
        PDF bytes
    """
    try:
        # Validate shipment
        if not shipment:
            raise ValueError("Shipment is required")

        # Get company info
        company_info = get_company_info()

        # Get status display name for both languages
        from .email_service import get_status_display_name

        try:
            status_display_ar = get_status_display_name(
                shipment.status, shipment.direction
            )
        except Exception as e:
            logger.warning(f"Error getting status display name: {str(e)}")
            status_display_ar = shipment.status or "Unknown"
        # For English, we'll use the same text for now (can be improved later)
        status_display_en = status_display_ar

        # Generate QR Code for tracking
        try:
            site_url = company_info.get("site_url", "https://medo-freight.eu")
            if not site_url.startswith("http"):
                site_url = f"https://{site_url}"
            tracking_url = f"{site_url}/tracking?shipment_id={shipment.id}"
        except Exception as e:
            logger.warning(f"Error generating tracking URL: {str(e)}")
            tracking_url = f"https://medo-freight.eu/tracking?shipment_id={shipment.id}"

        qr_code_base64 = None
        try:
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=4,
            )
            qr.add_data(tracking_url)
            qr.make(fit=True)

            # Create QR code image
            img = qr.make_image(fill_color="black", back_color="white")

            # Convert to base64
            img_buffer = io.BytesIO()
            img.save(img_buffer, format="PNG")
            img_buffer.seek(0)
            qr_code_base64 = base64.b64encode(img_buffer.read()).decode("utf-8")
        except Exception as e:
            logger.warning(f"Could not generate QR code: {str(e)}")
            qr_code_base64 = None

        # Get signature if exists (use invoice_signature for company signature)
        signature_base64 = None
        if shipment.invoice_signature:
            try:
                signature_path = shipment.invoice_signature.path
                if os.path.exists(signature_path):
                    with open(signature_path, "rb") as f:
                        signature_data = f.read()
                        signature_base64 = base64.b64encode(signature_data).decode(
                            "utf-8"
                        )
                        logger.info(
                            f"Successfully loaded signature for receipt {shipment.id}"
                        )
            except Exception as e:
                logger.warning(f"Could not read signature file: {str(e)}")

        # Generate barcode for receipt number
        receipt_number = shipment.shipment_number or str(shipment.id)
        barcode_base64 = generate_barcode(receipt_number)

        # Ensure receipt_number is never None
        display_receipt_number = shipment.shipment_number or f"LCL-{shipment.id:06d}"

        # Calculate shipment totals and prepare parcel items for receipt
        total_weight = 0.0
        chargeable_weight = 0.0
        total_pieces = 0
        receipt_items = []

        # Use calculate_invoice_totals to get proper product names and HS codes
        pricing = calculate_invoice_totals(shipment)

        parcels_data = shipment.parcels if shipment.parcels else []
        for idx, parcel_data in enumerate(parcels_data):
            if not parcel_data or not isinstance(parcel_data, dict):
                continue

            try:
                weight = float(parcel_data.get("weight", 0) or 0)
                cbm = float(parcel_data.get("cbm", 0) or 0)
                length = float(parcel_data.get("length", 0) or 0)
                width = float(parcel_data.get("width", 0) or 0)
                height = float(parcel_data.get("height", 0) or 0)
                repeat_count = int(parcel_data.get("repeatCount", 1) or 1)

                # Calculate volumetric weight: (L × W × H) / 6,000 or CBM × 167
                if length > 0 and width > 0 and height > 0:
                    volumetric_weight = (length * width * height) / 6000
                elif cbm > 0:
                    volumetric_weight = cbm * 167
                else:
                    volumetric_weight = 0

                # Chargeable weight = max(actual_weight, volumetric_weight)
                parcel_chargeable = max(weight, volumetric_weight)

                # Add to totals (multiply by repeat_count)
                total_weight += weight * repeat_count
                chargeable_weight += parcel_chargeable * repeat_count
                total_pieces += repeat_count

                # Get product name and HS code from pricing calculations if available
                product_name_ar = None
                product_name_en = None
                hs_code = parcel_data.get("hs_code") or parcel_data.get("hsCode")

                # Try to find matching calculation from pricing
                if idx < len(pricing.get("parcel_calculations", [])):
                    calc = pricing["parcel_calculations"][idx]
                    product_name_ar = calc.get("product_name_ar")
                    product_name_en = calc.get("product_name_en")
                    if not hs_code:
                        hs_code = calc.get("hs_code")

                # Fallback to parcel data
                if not product_name_ar and not product_name_en:
                    product_name_ar = parcel_data.get("description") or parcel_data.get(
                        "productName"
                    )
                    product_name_en = product_name_ar

                receipt_items.append(
                    {
                        "product_name_ar": product_name_ar or "General Goods",
                        "product_name_en": product_name_en or "General Goods",
                        "hs_code": hs_code or "N/A",
                        "quantity": repeat_count,
                        "description": parcel_data.get("description")
                        or parcel_data.get("notes"),
                    }
                )
            except (TypeError, ValueError) as e:
                logger.warning(f"Error calculating parcel totals: {str(e)}")
                continue

        # Prepare context for template
        context = {
            "shipment": shipment,
            "company": company_info,
            "language": language,
            "receipt_date": timezone.now(),
            "receipt_number": display_receipt_number,
            "status_display": {
                "ar": status_display_ar,
                "en": status_display_en,
            },
            "qr_code_base64": qr_code_base64,
            "barcode_base64": barcode_base64,
            "status_display_text": (
                status_display_ar if language == "ar" else status_display_en
            ),
            "tracking_url": tracking_url,
            "signature_base64": signature_base64,
            "total_weight": round(total_weight, 2),
            "chargeable_weight": round(chargeable_weight, 2),
            "total_pieces": total_pieces,
            "receipt_items": receipt_items,
        }

        # Render HTML template
        html_string = render_to_string("documents/receipt.html", context)

        # Generate PDF using WeasyPrint
        font_config = FontConfiguration()
        html = HTML(string=html_string, base_url=settings.BASE_DIR)

        # Generate PDF
        pdf_bytes = html.write_pdf(font_config=font_config)

        logger.info(f"Successfully generated receipt PDF for shipment {shipment.id}")
        return pdf_bytes

    except Exception as e:
        logger.error(f"Error generating receipt: {str(e)}", exc_info=True)
        raise


def save_receipt_to_storage(shipment: LCLShipment, pdf_bytes: bytes) -> str:
    """
    Save receipt PDF to storage and update shipment record.

    Args:
        shipment: LCLShipment instance
        pdf_bytes: PDF file bytes

    Returns:
        File path
    """
    try:
        from django.core.files.base import ContentFile

        # Generate filename
        filename = f"Receipt-{shipment.shipment_number}.pdf"

        # Save to FileField
        shipment.receipt_file.save(filename, ContentFile(pdf_bytes), save=False)
        shipment.receipt_generated_at = timezone.now()
        shipment.save()

        logger.info(
            f"Saved receipt to {shipment.receipt_file.path} for shipment {shipment.id}"
        )
        return shipment.receipt_file.path

    except Exception as e:
        logger.error(f"Error saving receipt to storage: {str(e)}", exc_info=True)
        raise


# ============================================================================
# Word Document Generation Functions
# ============================================================================


def _add_image_to_docx(doc, image_base64, width_inches=1.8):
    """Add base64 image to Word document"""
    try:
        if not image_base64:
            return None
        image_data = base64.b64decode(image_base64)
        image_stream = io.BytesIO(image_data)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(width_inches))
        return paragraph
    except Exception as e:
        logger.warning(f"Could not add image to Word document: {str(e)}")
        return None


def _set_cell_shading(cell, fill_color):
    """Set cell background color/shading"""
    try:
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_color)
        tcPr.append(shd)
    except Exception as e:
        logger.warning(f"Could not set cell shading: {str(e)}")


def _add_table_row(doc, cells_data, header=False, widths=None):
    """Add a table row to Word document"""
    if not hasattr(doc, "_current_table"):
        # Create new table
        table = doc.add_table(rows=1, cols=len(cells_data))
        table.style = "Light Grid Accent 1"
        if widths:
            for i, width in enumerate(widths):
                if i < len(table.columns):
                    table.columns[i].width = Inches(width)
        doc._current_table = table
    else:
        # Add row to existing table
        table = doc._current_table
        table.add_row()

    for i, cell_data in enumerate(cells_data):
        if i < len(table.rows[-1].cells):
            cell = table.rows[-1].cells[i]
            cell.text = str(cell_data) if cell_data else ""
            if header:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Dark background for header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                _set_cell_shading(cell, "1F2937")
            else:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    return table


def generate_packing_list_word(shipment: LCLShipment, language: str = "en") -> bytes:
    """
    Generate Packing List Word document for LCL shipment.

    Args:
        shipment: LCLShipment instance
        language: kept for future use (currently template is EN)

    Returns:
        Word document bytes (.docx)
    """
    try:
        # Reuse invoice pricing calculations
        pricing = calculate_invoice_totals(shipment)
        company_info = get_company_info()

        # Create Word document
        doc = Document()
        doc.sections[0].page_width = Inches(8.5)
        doc.sections[0].page_height = Inches(11)
        # Increased margins to prevent content overflow
        doc.sections[0].left_margin = Inches(0.5)
        doc.sections[0].right_margin = Inches(0.5)
        doc.sections[0].top_margin = Inches(0.5)
        doc.sections[0].bottom_margin = Inches(0.5)

        # Header section
        header_table = doc.add_table(rows=1, cols=2)
        header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Adjusted widths to fit within page (usable width: 8.5 - 1.0 = 7.5 inches)
        header_table.columns[0].width = Inches(4.0)
        header_table.columns[1].width = Inches(3.5)

        # Left column - Logo and company info
        left_cell = header_table.rows[0].cells[0]
        if company_info.get("logo_base64"):
            _add_image_to_docx(left_cell, company_info["logo_base64"], width_inches=1.8)

        p = left_cell.add_paragraph()
        run = p.add_run("MEDO-FREIGHT.EU SHIP. ROUTE. DELIVER")
        run.bold = True
        run.font.size = Pt(18)

        p = left_cell.add_paragraph()
        run = p.add_run("Logistics Service Provider")
        run.bold = True
        run.font.size = Pt(12)

        p = left_cell.add_paragraph()
        run = p.add_run(
            "Medo Freight\nTitanlaan 1, 4624 AX Bergen op Zoom, The Netherlands\nKvk nr: 75251663\nTAX nr: NL002518102B41\nEORI number: NL1320963189\nTel: +31 6 39 788 989  E-mail: contact@medo-freight.eu\nWebsite: http://medo-freight.eu"
        )
        run.font.size = Pt(11)

        # Right column - Barcode
        right_cell = header_table.rows[0].cells[1]
        tracking_url = f"{company_info['site_url']}/tracking"
        barcode_base64 = generate_tracking_barcode(tracking_url)
        if barcode_base64:
            _add_image_to_docx(right_cell, barcode_base64, width_inches=2.8)

        p = right_cell.add_paragraph()
        run = p.add_run("Middle East Office / مكتب الشرق الأوسط")
        run.bold = True
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = right_cell.add_paragraph()
        run = p.add_run(
            "Al Ikram Trading Co. / شركة الإكرام التجارية\nالرامسة (بجانب كراج البولمان) – الشرق الأوسط\nالمدينة الصناعية، الشيخ نجار (منطقة مكاتب الشحن الدولي)\nTel: +963 995 477 8188\nEmail: alikramtrading.co@gmail.com"
        )
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add horizontal line
        p = doc.add_paragraph()
        p.add_run("_" * 100)

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "Consolidated Export Packing List – Mixed Shipment (Personal & Commercial Goods)"
        )
        run.bold = True
        run.font.size = Pt(20)

        # Info grid table
        info_table = doc.add_table(rows=1, cols=5)
        info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Adjusted widths to fit within page (total: 7.5 inches)
        info_table.columns[0].width = Inches(1.5)
        info_table.columns[1].width = Inches(1.5)
        info_table.columns[2].width = Inches(1.5)
        info_table.columns[3].width = Inches(1.5)
        info_table.columns[4].width = Inches(1.5)

        info_cells = [
            "Invoice No.",
            "Invoice Date",
            "GST Reg No.",
            "Your Order No.",
            "Page No.",
        ]
        info_values = ["", "", "0.00%", "", "1"]
        for i, (label, value) in enumerate(zip(info_cells, info_values)):
            cell = info_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(14)
            run.bold = True

        # Address grid: For Account of Consignee / Deliver To Consignee
        address_table = doc.add_table(rows=1, cols=2)
        address_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        address_table.columns[0].width = Inches(3.75)
        address_table.columns[1].width = Inches(3.75)

        for i, label in enumerate(["For Account of Consignee", "Deliver To Consignee"]):
            cell = address_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            # Empty value field
            run = p.add_run("")
            run.font.size = Pt(12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Customer No. / Currency table
        customer_table = doc.add_table(rows=1, cols=2)
        customer_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        customer_table.columns[0].width = Inches(3.75)
        customer_table.columns[1].width = Inches(3.75)

        customer_labels = ["Customer No.", "Currency"]
        customer_values = ["", "EUR"]
        for i, (label, value) in enumerate(zip(customer_labels, customer_values)):
            cell = customer_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(18)
            run.bold = True

        # Payment Terms / Sales Area table
        payment_table = doc.add_table(rows=1, cols=2)
        payment_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        payment_table.columns[0].width = Inches(3.75)
        payment_table.columns[1].width = Inches(3.75)

        payment_labels = ["Payment Terms", "Sales Area"]
        payment_values = ["", ""]
        for i, (label, value) in enumerate(zip(payment_labels, payment_values)):
            cell = payment_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(18)
            run.bold = True

        # Incoterms / Bank Details table
        incoterms_table = doc.add_table(rows=1, cols=2)
        incoterms_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        incoterms_table.columns[0].width = Inches(3.75)
        incoterms_table.columns[1].width = Inches(3.75)

        incoterms_labels = ["Incoterms", "Bank Details"]
        incoterms_values = ["", ""]
        for i, (label, value) in enumerate(zip(incoterms_labels, incoterms_values)):
            cell = incoterms_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(18)
            run.bold = True

        # Shipping Details Grid (11 boxes in 2 rows)
        shipping_labels = [
            "Bill of Lading",
            "Loading Port / City",
            "Destination Port",
            "Despatch",
            "Mode / Vessel Name",
            "Container No.",
            "Voyage No.",
            "ETD",
            "ETA",
            "Shipping",
            "Seal No.",
        ]
        shipping_values = ["", "", "", "", "", "", "", "", "", "LCL", ""]

        # Create shipping details table with 6 columns (will wrap to next row)
        shipping_table = doc.add_table(rows=2, cols=6)
        shipping_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i in range(6):
            shipping_table.columns[i].width = Inches(1.25)

        for idx, (label, value) in enumerate(zip(shipping_labels, shipping_values)):
            row_idx = idx // 6
            col_idx = idx % 6
            if row_idx < 2 and col_idx < 6:
                cell = shipping_table.rows[row_idx].cells[col_idx]
                p = cell.add_paragraph()
                run = p.add_run(label)
                run.font.size = Pt(9)
                run.bold = True
                p = cell.add_paragraph()
                run = p.add_run(value)
                run.font.size = Pt(12)
                run.bold = True

        # Aggregate totals
        total_cbm = 0.0
        total_packages = 0
        total_weight = 0.0

        for item in pricing.get("parcel_calculations", []):
            try:
                cbm_value = float(item.get("cbm", 0) or 0)
            except (TypeError, ValueError):
                cbm_value = 0.0
            total_cbm += cbm_value

            try:
                weight_value = float(item.get("weight", 0) or 0)
            except (TypeError, ValueError):
                weight_value = 0.0
            repeat_count = int(item.get("repeat_count", 1) or 1)
            total_weight += weight_value * repeat_count
            total_packages += repeat_count

        # Product table - adjusted widths to fit within 7.5 inches
        product_table = doc.add_table(rows=1, cols=8)
        product_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        product_table.columns[0].width = Inches(0.75)  # Product Code
        product_table.columns[1].width = Inches(1.4)  # Client Name
        product_table.columns[2].width = Inches(1.8)  # Description
        product_table.columns[3].width = Inches(0.9)  # Shipment Type
        product_table.columns[4].width = Inches(0.75)  # Origin
        product_table.columns[5].width = Inches(0.75)  # HS Code
        product_table.columns[6].width = Inches(0.6)  # CBM
        product_table.columns[7].width = Inches(0.55)  # Unit

        # Header row
        headers = [
            "Product Code",
            "Client Name",
            "Description",
            "Shipment Type",
            "Origin",
            "HS Code",
            "CBM",
            "Unit",
        ]
        for i, header in enumerate(headers):
            cell = product_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Dark background
            _set_cell_shading(cell, "1F2937")
            run.font.color.rgb = RGBColor(255, 255, 255)

        # Data rows
        for item in pricing.get("parcel_calculations", []):
            row = product_table.add_row()
            row.cells[0].text = shipment.shipment_number or ""
            row.cells[1].text = shipment.sender_name or ""

            desc_text = item.get("product_name_en", "")
            if item.get("product_name_ar"):
                desc_text += f"\n{item['product_name_ar']}"
            row.cells[2].text = desc_text

            shipment_type = item.get("shipment_type") or shipment.shipment_type
            if shipment_type == "personal":
                row.cells[3].text = "Personal / شخصي"
            elif shipment_type == "commercial":
                row.cells[3].text = "Commercial / تجاري"
            else:
                row.cells[3].text = "-"

            row.cells[4].text = ""
            row.cells[5].text = item.get("hs_code", "") or ""
            row.cells[6].text = f"{item.get('cbm', 0):.3f}"
            row.cells[7].text = "PCS"

            # Center align some columns
            for i in [3, 4, 5, 6, 7]:
                row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Second product table - adjusted widths to fit within 7.5 inches
        product_table2 = doc.add_table(rows=1, cols=3)
        product_table2.alignment = WD_TABLE_ALIGNMENT.LEFT
        product_table2.columns[0].width = Inches(2.5)  # Qty
        product_table2.columns[1].width = Inches(2.5)  # Gross Wt
        product_table2.columns[2].width = Inches(2.5)  # Nett Wt

        headers2 = ["Qty", "Gross Wt", "Nett Wt"]
        for i, header in enumerate(headers2):
            cell = product_table2.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell, "1F2937")
            run.font.color.rgb = RGBColor(255, 255, 255)

        for item in pricing.get("parcel_calculations", []):
            row = product_table2.add_row()
            row.cells[0].text = str(item.get("repeat_count", 1))
            row.cells[1].text = f"{item.get('weight', 0):.2f}"
            row.cells[2].text = ""

            # Right align weight columns
            for i in [1, 2]:
                row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Totals table - adjusted widths to fit within 7.5 inches
        totals_table = doc.add_table(rows=1, cols=1)
        totals_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        totals_table.columns[0].width = Inches(7.5)

        totals_labels = ["PACKAGES"]
        totals_values = [str(total_packages)]

        for i, (label, value) in enumerate(zip(totals_labels, totals_values)):
            cell = totals_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(18)
            run.bold = True

        # Disclaimer
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "This packing list is issued for cargo identification purposes only.\n\n"
            "All packages have been checked, weighed, and sealed before loading"
        )
        run.font.size = Pt(11)
        run.bold = True
        p = doc.add_paragraph()

        # Footer notes
        p = doc.add_paragraph()
        run = p.add_run("A - Legal & Customs Declaration\n")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(
            "MEDO-B2B EU acts solely as a freight consolidator and export agent on behalf of multiple clients. "
            "The declared values are for customs declaration purposes only and do not represent sales or transfer of ownership. "
            "This shipment is handled under EU Export Compliance (EX-A). MEDO-B2B EU bears no ownership or sales relation "
            "to the goods. All clients have submitted their own declaration forms confirming that their items are personal "
            "or commercial as described.\n\n"
        )

        run = p.add_run("B - Freight & Service Terms\n")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(
            "Payment is due only for freight and handling services, by bank transfer or cash, 100% prepaid before departure. "
            "The consignee acknowledges that all goods were inspected and accepted before loading. No return, refund, or "
            "complaint is accepted after the container has been sealed and departed. MEDO-B2B EU provides transport "
            "services under CIF Incoterms (Cost, Insurance, Freight) unless otherwise agreed in writing. Our general "
            "terms and conditions apply to all shipments and can be provided upon request."
        )

        # Signature section - adjusted widths to fit within 7.5 inches
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        sig_table.columns[0].width = Inches(3.75)
        sig_table.columns[1].width = Inches(3.75)

        for i in range(2):
            cell = sig_table.rows[0].cells[i]
            p = cell.add_paragraph()
            if i == 0:
                run = p.add_run("Authorized Signature & Stamp")
            else:
                run = p.add_run("Date")
            run.bold = True
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add space for signature
            for _ in range(5):
                cell.add_paragraph()

            p = cell.add_paragraph()
            p.add_run("_" * 30)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        logger.info(
            f"Successfully generated packing list Word document for shipment {shipment.id}"
        )
        return doc_bytes.getvalue()

    except Exception as e:
        logger.error(
            f"Error generating packing list Word document: {str(e)}", exc_info=True
        )
        raise


def generate_consolidated_export_invoice_word(
    shipment: LCLShipment, language: str = "en"
) -> bytes:
    """
    Generate Consolidated Export Invoice Word document for LCL shipment.

    Args:
        shipment: LCLShipment instance
        language: kept for future use (currently template is EN)

    Returns:
        Word document bytes (.docx)
    """
    try:
        # Reuse invoice pricing calculations
        pricing = calculate_invoice_totals(shipment)
        company_info = get_company_info()

        # Create Word document
        doc = Document()
        doc.sections[0].page_width = Inches(8.5)
        doc.sections[0].page_height = Inches(11)
        # Increased margins to prevent content overflow
        doc.sections[0].left_margin = Inches(0.5)
        doc.sections[0].right_margin = Inches(0.5)
        doc.sections[0].top_margin = Inches(0.5)
        doc.sections[0].bottom_margin = Inches(0.5)

        # Header section (same as packing list)
        header_table = doc.add_table(rows=1, cols=2)
        # Adjusted widths to fit within page (usable width: 8.5 - 1.0 = 7.5 inches)
        header_table.columns[0].width = Inches(4.0)
        header_table.columns[1].width = Inches(3.5)

        left_cell = header_table.rows[0].cells[0]
        if company_info.get("logo_base64"):
            _add_image_to_docx(left_cell, company_info["logo_base64"], width_inches=1.8)

        p = left_cell.add_paragraph()
        run = p.add_run("MEDO-FREIGHT.EU SHIP. ROUTE. DELIVER")
        run.bold = True
        run.font.size = Pt(18)

        p = left_cell.add_paragraph()
        run = p.add_run("Logistics Service Provider")
        run.bold = True
        run.font.size = Pt(12)

        p = left_cell.add_paragraph()
        run = p.add_run(
            "Medo Freight\nTitanlaan 1, 4624 AX Bergen op Zoom, The Netherlands\nKvk nr: 75251663\nTAX nr: NL002518102B41\nEORI number: NL1320963189\nTel: +31 6 39 788 989  E-mail: contact@medo-freight.eu\nWebsite: http://medo-freight.eu"
        )
        run.font.size = Pt(11)

        right_cell = header_table.rows[0].cells[1]
        tracking_url = f"{company_info['site_url']}/tracking"
        barcode_base64 = generate_tracking_barcode(tracking_url)
        if barcode_base64:
            _add_image_to_docx(right_cell, barcode_base64, width_inches=2.8)

        p = right_cell.add_paragraph()
        run = p.add_run("Middle East Office / مكتب الشرق الأوسط")
        run.bold = True
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = right_cell.add_paragraph()
        run = p.add_run(
            "Al Ikram Trading Co. / شركة الإكرام التجارية\nالرامسة (بجانب كراج البولمان) – الشرق الأوسط\nالمدينة الصناعية، الشيخ نجار (منطقة مكاتب الشحن الدولي)\nTel: +963 995 477 8188\nEmail: alikramtrading.co@gmail.com"
        )
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = doc.add_paragraph()
        p.add_run("_" * 100)

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "Consolidated Export Invoice – Mixed Shipment (Personal & Commercial Goods)"
        )
        run.bold = True
        run.font.size = Pt(20)

        # Info grid table
        info_table = doc.add_table(rows=1, cols=5)
        # Adjusted widths to fit within page (total: 7.5 inches)
        for i in range(5):
            info_table.columns[i].width = Inches(1.5)

        info_cells = [
            "Invoice No.",
            "Invoice Date",
            "GST Reg No.",
            "Your Order No.",
            "Page No.",
        ]
        info_values = ["", "", "0.00%", "", "1"]
        for i, (label, value) in enumerate(zip(info_cells, info_values)):
            cell = info_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(14)
            run.bold = True

        # Address grid: For Account of Consignee / Deliver To Consignee
        address_table = doc.add_table(rows=1, cols=2)
        address_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        address_table.columns[0].width = Inches(3.75)
        address_table.columns[1].width = Inches(3.75)

        for i, label in enumerate(["For Account of Consignee", "Deliver To Consignee"]):
            cell = address_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            # Empty value field
            run = p.add_run("")
            run.font.size = Pt(12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Customer No. / Currency table
        customer_table = doc.add_table(rows=1, cols=2)
        customer_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        customer_table.columns[0].width = Inches(3.75)
        customer_table.columns[1].width = Inches(3.75)

        customer_labels = ["Customer No.", "Currency"]
        customer_values = ["", "EUR"]
        for i, (label, value) in enumerate(zip(customer_labels, customer_values)):
            cell = customer_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(18)
            run.bold = True

        # Payment Terms / Sales Area table
        payment_table = doc.add_table(rows=1, cols=2)
        payment_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        payment_table.columns[0].width = Inches(3.75)
        payment_table.columns[1].width = Inches(3.75)

        payment_labels = ["Payment Terms", "Sales Area"]
        payment_values = ["", ""]
        for i, (label, value) in enumerate(zip(payment_labels, payment_values)):
            cell = payment_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(18)
            run.bold = True

        # Incoterms / Bank Details table
        incoterms_table = doc.add_table(rows=1, cols=2)
        incoterms_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        incoterms_table.columns[0].width = Inches(3.75)
        incoterms_table.columns[1].width = Inches(3.75)

        incoterms_labels = ["Incoterms", "Bank Details"]
        incoterms_values = ["", ""]
        for i, (label, value) in enumerate(zip(incoterms_labels, incoterms_values)):
            cell = incoterms_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(18)
            run.bold = True

        # Shipping Details Grid (11 boxes in 2 rows)
        shipping_labels = [
            "Bill of Lading",
            "Loading Port / City",
            "Destination Port",
            "Despatch",
            "Mode / Vessel Name",
            "Container No.",
            "Voyage No.",
            "ETD",
            "ETA",
            "Shipping",
            "Seal No.",
        ]
        shipping_values = ["", "", "", "", "", "", "", "", "", "LCL", ""]

        # Create shipping details table with 6 columns (will wrap to next row)
        shipping_table = doc.add_table(rows=2, cols=6)
        shipping_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i in range(6):
            shipping_table.columns[i].width = Inches(1.25)

        for idx, (label, value) in enumerate(zip(shipping_labels, shipping_values)):
            row_idx = idx // 6
            col_idx = idx % 6
            if row_idx < 2 and col_idx < 6:
                cell = shipping_table.rows[row_idx].cells[col_idx]
                p = cell.add_paragraph()
                run = p.add_run(label)
                run.font.size = Pt(9)
                run.bold = True
                p = cell.add_paragraph()
                run = p.add_run(value)
                run.font.size = Pt(12)
                run.bold = True

        # Aggregate totals
        total_cbm = 0.0
        total_packages = 0
        total_weight = 0.0
        shipment_types = []

        for item in pricing.get("parcel_calculations", []):
            try:
                cbm_value = float(item.get("cbm", 0) or 0)
            except (TypeError, ValueError):
                cbm_value = 0.0
            total_cbm += cbm_value

            try:
                weight_value = float(item.get("weight", 0) or 0)
            except (TypeError, ValueError):
                weight_value = 0.0
            repeat_count = int(item.get("repeat_count", 1) or 1)
            total_weight += weight_value * repeat_count
            total_packages += repeat_count

            if shipment.parcels:
                for parcel in shipment.parcels:
                    if isinstance(parcel, dict):
                        parcel_type = parcel.get("shipmentType") or parcel.get(
                            "shipment_type"
                        )
                        if parcel_type:
                            shipment_types.append(parcel_type)

        unique_types = list(set(shipment_types))
        is_personal_only = len(unique_types) == 1 and unique_types[0] == "personal"
        is_commercial_only = len(unique_types) == 1 and unique_types[0] == "commercial"
        is_mixed = len(unique_types) > 1

        # Product table (same structure as packing list) - adjusted widths to fit within 7.5 inches
        product_table = doc.add_table(rows=1, cols=8)
        product_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        product_table.columns[0].width = Inches(0.75)  # Product Code
        product_table.columns[1].width = Inches(1.4)  # Client Name
        product_table.columns[2].width = Inches(1.8)  # Description
        product_table.columns[3].width = Inches(0.9)  # Shipment Type
        product_table.columns[4].width = Inches(0.75)  # Origin
        product_table.columns[5].width = Inches(0.75)  # HS Code
        product_table.columns[6].width = Inches(0.6)  # CBM
        product_table.columns[7].width = Inches(0.55)  # Unit

        headers = [
            "Product Code",
            "Client Name",
            "Description",
            "Shipment Type",
            "Origin",
            "HS Code",
            "CBM",
            "Unit",
        ]
        for i, header in enumerate(headers):
            cell = product_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell, "1F2937")
            run.font.color.rgb = RGBColor(255, 255, 255)

        for item in pricing.get("parcel_calculations", []):
            row = product_table.add_row()
            row.cells[0].text = shipment.shipment_number or ""
            row.cells[1].text = shipment.sender_name or ""

            desc_text = item.get("product_name_en", "")
            if item.get("product_name_ar"):
                desc_text += f"\n{item['product_name_ar']}"
            row.cells[2].text = desc_text

            shipment_type = item.get("shipment_type") or shipment.shipment_type
            if shipment_type == "personal":
                row.cells[3].text = "Personal / شخصي"
            elif shipment_type == "commercial":
                row.cells[3].text = "Commercial / تجاري"
            else:
                row.cells[3].text = "-"

            row.cells[4].text = ""
            row.cells[5].text = item.get("hs_code", "") or ""
            row.cells[6].text = f"{item.get('cbm', 0):.3f}"
            row.cells[7].text = "PCS"

            for i in [3, 4, 5, 6, 7]:
                row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Second product table - adjusted widths to fit within 7.5 inches
        product_table2 = doc.add_table(rows=1, cols=5)
        product_table2.alignment = WD_TABLE_ALIGNMENT.LEFT
        product_table2.columns[0].width = Inches(1.3)  # Qty
        product_table2.columns[1].width = Inches(1.5)  # Unit Price
        product_table2.columns[2].width = Inches(1.5)  # Value
        product_table2.columns[3].width = Inches(1.5)  # Gross Wt
        product_table2.columns[4].width = Inches(1.7)  # Nett Wt

        headers2 = ["Qty", "Unit Price", "Value", "Gross Wt", "Nett Wt"]
        for i, header in enumerate(headers2):
            cell = product_table2.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell, "1F2937")
            run.font.color.rgb = RGBColor(255, 255, 255)

        for item in pricing.get("parcel_calculations", []):
            row = product_table2.add_row()
            row.cells[0].text = str(item.get("repeat_count", 1))

            shipment_type = item.get("shipment_type") or shipment.shipment_type
            if shipment_type == "commercial":
                row.cells[1].text = ""
                row.cells[2].text = ""
            elif shipment_type == "personal":
                row.cells[1].text = "€0.00"
                row.cells[2].text = "€0.00"
            else:
                if item.get("price_per_kg"):
                    row.cells[1].text = f"€{item['price_per_kg']:.2f}"
                row.cells[2].text = f"€{item.get('price_by_weight', 0):.2f}"

            row.cells[3].text = f"{item.get('weight', 0):.2f}"
            row.cells[4].text = ""

            for i in [1, 2, 3, 4]:
                row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Totals table - adjusted widths to fit within 7.5 inches
        totals_table = doc.add_table(rows=1, cols=3)
        totals_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        totals_table.columns[0].width = Inches(2.5)
        totals_table.columns[1].width = Inches(2.5)
        totals_table.columns[2].width = Inches(2.5)

        totals_labels = ["SALES TAX", "PACKAGES", "TOTAL INVOICE VALUE"]
        totals_values = ["", str(total_packages), f"EUR {pricing['total_price']:.2f}"]

        for i, (label, value) in enumerate(zip(totals_labels, totals_values)):
            cell = totals_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(18)
            run.bold = True

        # Footer notes
        p = doc.add_paragraph()
        run = p.add_run("A - Legal & Customs Declaration\n")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(
            "MEDO-B2B EU acts solely as a freight consolidator and export agent on behalf of multiple clients. "
            "The declared values are for customs declaration purposes only and do not represent sales or transfer of ownership. "
            "This shipment is handled under EU Export Compliance (EX-A). MEDO-B2B EU bears no ownership or sales relation "
            "to the goods. All clients have submitted their own declaration forms confirming that their items are personal "
            "or commercial as described.\n\n"
        )

        run = p.add_run("B - Freight & Service Terms\n")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(
            "Payment is due only for freight and handling services, by bank transfer or cash, 100% prepaid before departure. "
            "The consignee acknowledges that all goods were inspected and accepted before loading. No return, refund, or "
            "complaint is accepted after the container has been sealed and departed. MEDO-B2B EU provides transport "
            "services under CIF Incoterms (Cost, Insurance, Freight) unless otherwise agreed in writing. Our general "
            "terms and conditions apply to all shipments and can be provided upon request."
        )

        # Signature section - adjusted widths to fit within 7.5 inches
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        sig_table.columns[0].width = Inches(3.75)
        sig_table.columns[1].width = Inches(3.75)

        for i in range(2):
            cell = sig_table.rows[0].cells[i]
            p = cell.add_paragraph()
            if i == 0:
                run = p.add_run("Authorized Signature & Stamp")
            else:
                run = p.add_run("Date")
            run.bold = True
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for _ in range(5):
                cell.add_paragraph()

            p = cell.add_paragraph()
            p.add_run("_" * 30)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        logger.info(
            f"Successfully generated consolidated export invoice Word document for shipment {shipment.id}"
        )
        return doc_bytes.getvalue()

    except Exception as e:
        logger.error(
            f"Error generating consolidated export invoice Word document: {str(e)}",
            exc_info=True,
        )
        raise


def generate_consolidated_packing_list_word(
    shipments: List[LCLShipment], language: str = "en"
) -> bytes:
    """
    Generate Consolidated Packing List Word document for multiple LCL shipments.

    Args:
        shipments: List of LCLShipment instances
        language: kept for future use (currently template is EN)

    Returns:
        Word document bytes (.docx)
    """
    try:
        if not shipments:
            raise ValueError("No shipments provided for consolidated packing list")

        company_info = get_company_info()

        # Create Word document
        doc = Document()
        doc.sections[0].page_width = Inches(8.5)
        doc.sections[0].page_height = Inches(11)
        # Increased margins to prevent content overflow
        doc.sections[0].left_margin = Inches(0.5)
        doc.sections[0].right_margin = Inches(0.5)
        doc.sections[0].top_margin = Inches(0.5)
        doc.sections[0].bottom_margin = Inches(0.5)

        # Header (same as single shipment)
        header_table = doc.add_table(rows=1, cols=2)
        # Adjusted widths to fit within page (usable width: 8.5 - 1.0 = 7.5 inches)
        header_table.columns[0].width = Inches(4.0)
        header_table.columns[1].width = Inches(3.5)

        left_cell = header_table.rows[0].cells[0]
        if company_info.get("logo_base64"):
            _add_image_to_docx(left_cell, company_info["logo_base64"], width_inches=1.8)

        p = left_cell.add_paragraph()
        run = p.add_run("MEDO-FREIGHT.EU SHIP. ROUTE. DELIVER")
        run.bold = True
        run.font.size = Pt(18)

        p = left_cell.add_paragraph()
        run = p.add_run("Logistics Service Provider")
        run.bold = True
        run.font.size = Pt(12)

        p = left_cell.add_paragraph()
        run = p.add_run(
            "Medo Freight\nTitanlaan 1, 4624 AX Bergen op Zoom, The Netherlands\nKvk nr: 75251663\nTAX nr: NL002518102B41\nEORI number: NL1320963189\nTel: +31 6 39 788 989  E-mail: contact@medo-freight.eu\nWebsite: http://medo-freight.eu"
        )
        run.font.size = Pt(11)

        right_cell = header_table.rows[0].cells[1]
        tracking_url = f"{company_info['site_url']}/tracking"
        barcode_base64 = generate_tracking_barcode(tracking_url)
        if barcode_base64:
            _add_image_to_docx(right_cell, barcode_base64, width_inches=2.8)

        p = right_cell.add_paragraph()
        run = p.add_run("Middle East Office / مكتب الشرق الأوسط")
        run.bold = True
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = right_cell.add_paragraph()
        run = p.add_run(
            "Al Ikram Trading Co. / شركة الإكرام التجارية\nالرامسة (بجانب كراج البولمان) – الشرق الأوسط\nالمدينة الصناعية، الشيخ نجار (منطقة مكاتب الشحن الدولي)\nTel: +963 995 477 8188\nEmail: alikramtrading.co@gmail.com"
        )
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = doc.add_paragraph()
        p.add_run("_" * 100)

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "Consolidated Export Packing List – Mixed Shipment (Personal & Commercial Goods)"
        )
        run.bold = True
        run.font.size = Pt(20)

        # Info grid
        info_table = doc.add_table(rows=1, cols=5)
        for i in range(5):
            info_table.columns[i].width = Inches(1.7)

        info_cells = [
            "Invoice No.",
            "Invoice Date",
            "GST Reg No.",
            "Your Order No.",
            "Page No.",
        ]
        info_values = ["", "", "0.00%", "", "1"]
        for i, (label, value) in enumerate(zip(info_cells, info_values)):
            cell = info_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(14)
            run.bold = True

        # Aggregate data from all shipments
        grand_total_cbm = 0.0
        grand_total_packages = 0
        grand_total_weight = 0.0
        grand_total_value = 0.0
        all_shipment_data = []

        for shipment in shipments:
            pricing = calculate_invoice_totals(shipment)

            total_cbm = 0.0
            total_packages = 0
            total_weight = 0.0
            total_value = 0.0

            shipment_items = []
            for item in pricing.get("parcel_calculations", []):
                try:
                    cbm_value = float(item.get("cbm", 0) or 0)
                except (TypeError, ValueError):
                    cbm_value = 0.0
                total_cbm += cbm_value

                try:
                    weight_value = float(item.get("weight", 0) or 0)
                except (TypeError, ValueError):
                    weight_value = 0.0
                repeat_count = int(item.get("repeat_count", 1) or 1)
                total_weight += weight_value * repeat_count
                total_packages += repeat_count

                try:
                    price_value = float(item.get("price_by_weight", 0) or 0)
                except (TypeError, ValueError):
                    price_value = 0.0
                total_value += price_value

                shipment_items.append(
                    {
                        "shipment_number": shipment.shipment_number
                        or f"#{shipment.id}",
                        "product_name_en": item.get("product_name_en", ""),
                        "product_name_ar": item.get("product_name_ar", ""),
                        "hs_code": item.get("hs_code", ""),
                        "cbm": float(item.get("cbm", 0) or 0),
                        "repeat_count": int(item.get("repeat_count", 1) or 1),
                        "price_by_weight": float(item.get("price_by_weight", 0) or 0),
                        "weight": float(item.get("weight", 0) or 0),
                    }
                )

            grand_total_cbm += total_cbm
            grand_total_packages += total_packages
            grand_total_weight += total_weight
            grand_total_value += total_value

            all_shipment_data.append(
                {
                    "shipment": shipment,
                    "pricing": pricing,
                    "items": shipment_items,
                    "total_cbm": total_cbm,
                    "total_packages": total_packages,
                    "total_weight": total_weight,
                    "total_value": total_value,
                }
            )

        # Product table - adjusted widths to fit within 7.5 inches
        product_table = doc.add_table(rows=1, cols=8)
        product_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        product_table.columns[0].width = Inches(0.75)  # Product Code
        product_table.columns[1].width = Inches(1.4)  # Client Name
        product_table.columns[2].width = Inches(1.8)  # Description
        product_table.columns[3].width = Inches(0.9)  # Shipment Type
        product_table.columns[4].width = Inches(0.75)  # Origin
        product_table.columns[5].width = Inches(0.75)  # HS Code
        product_table.columns[6].width = Inches(0.6)  # CBM
        product_table.columns[7].width = Inches(0.55)  # Unit

        headers = [
            "Product Code",
            "Client Name",
            "Description",
            "Shipment Type",
            "Origin",
            "HS Code",
            "CBM",
            "Unit",
        ]
        for i, header in enumerate(headers):
            cell = product_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell, "1F2937")
            run.font.color.rgb = RGBColor(255, 255, 255)

        for shipment_info in all_shipment_data:
            for item in shipment_info["items"]:
                row = product_table.add_row()
                row.cells[0].text = item["shipment_number"]
                row.cells[1].text = shipment_info["shipment"].sender_name or ""

                desc_text = item["product_name_en"]
                if item["product_name_ar"]:
                    desc_text += f"\n{item['product_name_ar']}"
                row.cells[2].text = desc_text

                # Get shipment type from shipment
                shipment_type = shipment_info["shipment"].shipment_type
                if shipment_type == "personal":
                    row.cells[3].text = "Personal / شخصي"
                elif shipment_type == "commercial":
                    row.cells[3].text = "Commercial / تجاري"
                else:
                    row.cells[3].text = "-"

                row.cells[4].text = ""
                row.cells[5].text = item["hs_code"] or ""
                row.cells[6].text = f"{item['cbm']:.3f}"
                row.cells[7].text = "PCS"

                for i in [3, 4, 5, 6, 7]:
                    row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Second product table - adjusted widths to fit within 7.5 inches
        product_table2 = doc.add_table(rows=1, cols=3)
        product_table2.alignment = WD_TABLE_ALIGNMENT.LEFT
        product_table2.columns[0].width = Inches(2.5)  # Qty
        product_table2.columns[1].width = Inches(2.5)  # Gross Wt
        product_table2.columns[2].width = Inches(2.5)  # Nett Wt

        headers2 = ["Qty", "Gross Wt", "Nett Wt"]
        for i, header in enumerate(headers2):
            cell = product_table2.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell, "1F2937")
            run.font.color.rgb = RGBColor(255, 255, 255)

        for shipment_info in all_shipment_data:
            for item in shipment_info["items"]:
                row = product_table2.add_row()
                row.cells[0].text = str(item["repeat_count"])
                row.cells[1].text = f"{item['weight']:.2f}"
                row.cells[2].text = ""

                # Right align weight columns
                for i in [1, 2]:
                    row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Totals table - adjusted widths to fit within 7.5 inches
        totals_table = doc.add_table(rows=1, cols=1)
        totals_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        totals_table.columns[0].width = Inches(7.5)

        totals_labels = ["PACKAGES"]
        totals_values = [str(grand_total_packages)]

        for i, (label, value) in enumerate(zip(totals_labels, totals_values)):
            cell = totals_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(18)
            run.bold = True

        # Disclaimer
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "This packing list is issued for cargo identification purposes only.\n\n"
            "All packages have been checked, weighed, and sealed before loading"
        )
        run.font.size = Pt(11)
        run.bold = True
        p = doc.add_paragraph()

        # Footer notes
        p = doc.add_paragraph()
        run = p.add_run("A - Legal & Customs Declaration\n")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(
            "MEDO-B2B EU acts solely as a freight consolidator and export agent on behalf of multiple clients. "
            "The declared values are for customs declaration purposes only and do not represent sales or transfer of ownership. "
            "This shipment is handled under EU Export Compliance (EX-A). MEDO-B2B EU bears no ownership or sales relation "
            "to the goods. All clients have submitted their own declaration forms confirming that their items are personal "
            "or commercial as described.\n\n"
        )

        run = p.add_run("B - Freight & Service Terms\n")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(
            "Payment is due only for freight and handling services, by bank transfer or cash, 100% prepaid before departure. "
            "The consignee acknowledges that all goods were inspected and accepted before loading. No return, refund, or "
            "complaint is accepted after the container has been sealed and departed. MEDO-B2B EU provides transport "
            "services under CIF Incoterms (Cost, Insurance, Freight) unless otherwise agreed in writing. Our general "
            "terms and conditions apply to all shipments and can be provided upon request."
        )

        # Signature section - adjusted widths to fit within 7.5 inches
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        sig_table.columns[0].width = Inches(3.75)
        sig_table.columns[1].width = Inches(3.75)

        for i in range(2):
            cell = sig_table.rows[0].cells[i]
            p = cell.add_paragraph()
            if i == 0:
                run = p.add_run("Authorized Signature & Stamp")
            else:
                run = p.add_run("Date")
            run.bold = True
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for _ in range(5):
                cell.add_paragraph()

            p = cell.add_paragraph()
            p.add_run("_" * 30)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        logger.info(
            f"Successfully generated consolidated packing list Word document for {len(shipments)} shipments"
        )
        return doc_bytes.getvalue()

    except Exception as e:
        logger.error(
            f"Error generating consolidated packing list Word document: {str(e)}",
            exc_info=True,
        )
        raise


def generate_consolidated_export_invoice_bulk_word(
    shipments: List[LCLShipment],
    language: str = "en",
    invoice_number: Optional[str] = None,
) -> bytes:
    """
    Generate Consolidated Export Invoice Word document for multiple LCL shipments.

    Args:
        shipments: List of LCLShipment instances
        language: kept for future use (currently template is EN)
        invoice_number: Optional invoice number

    Returns:
        Word document bytes (.docx)
    """
    try:
        if not shipments:
            raise ValueError("No shipments provided for consolidated export invoice")

        company_info = get_company_info()

        # Create Word document
        doc = Document()
        doc.sections[0].page_width = Inches(8.5)
        doc.sections[0].page_height = Inches(11)
        # Increased margins to prevent content overflow
        doc.sections[0].left_margin = Inches(0.5)
        doc.sections[0].right_margin = Inches(0.5)
        doc.sections[0].top_margin = Inches(0.5)
        doc.sections[0].bottom_margin = Inches(0.5)

        # Header (same structure)
        header_table = doc.add_table(rows=1, cols=2)
        # Adjusted widths to fit within page (usable width: 8.5 - 1.0 = 7.5 inches)
        header_table.columns[0].width = Inches(4.0)
        header_table.columns[1].width = Inches(3.5)

        left_cell = header_table.rows[0].cells[0]
        if company_info.get("logo_base64"):
            _add_image_to_docx(left_cell, company_info["logo_base64"], width_inches=1.8)

        p = left_cell.add_paragraph()
        run = p.add_run("MEDO-FREIGHT.EU SHIP. ROUTE. DELIVER")
        run.bold = True
        run.font.size = Pt(18)

        p = left_cell.add_paragraph()
        run = p.add_run("Logistics Service Provider")
        run.bold = True
        run.font.size = Pt(12)

        p = left_cell.add_paragraph()
        run = p.add_run(
            "Medo Freight\nTitanlaan 1, 4624 AX Bergen op Zoom, The Netherlands\nKvk nr: 75251663\nTAX nr: NL002518102B41\nEORI number: NL1320963189\nTel: +31 6 39 788 989  E-mail: contact@medo-freight.eu\nWebsite: http://medo-freight.eu"
        )
        run.font.size = Pt(11)

        right_cell = header_table.rows[0].cells[1]
        tracking_url = f"{company_info['site_url']}/tracking"
        barcode_base64 = generate_tracking_barcode(tracking_url)
        if barcode_base64:
            _add_image_to_docx(right_cell, barcode_base64, width_inches=2.8)

        p = right_cell.add_paragraph()
        run = p.add_run("Middle East Office / مكتب الشرق الأوسط")
        run.bold = True
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = right_cell.add_paragraph()
        run = p.add_run(
            "Al Ikram Trading Co. / شركة الإكرام التجارية\nالرامسة (بجانب كراج البولمان) – الشرق الأوسط\nالمدينة الصناعية، الشيخ نجار (منطقة مكاتب الشحن الدولي)\nTel: +963 995 477 8188\nEmail: alikramtrading.co@gmail.com"
        )
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = doc.add_paragraph()
        p.add_run("_" * 100)

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "Consolidated Export Invoice – Mixed Shipment (Personal & Commercial Goods)"
        )
        run.bold = True
        run.font.size = Pt(20)

        # Info grid
        info_table = doc.add_table(rows=1, cols=5)
        info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Adjusted widths to fit within page (total: 7.5 inches)
        for i in range(5):
            info_table.columns[i].width = Inches(1.5)

        invoice_num = (
            invoice_number
            or f"INV-{timezone.now().strftime('%Y%m%d')}-{len(shipments)}"
        )
        info_cells = [
            "Invoice No.",
            "Invoice Date",
            "GST Reg No.",
            "Your Order No.",
            "Page No.",
        ]
        info_values = [invoice_num, "", "0.00%", "", "1"]
        for i, (label, value) in enumerate(zip(info_cells, info_values)):
            cell = info_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(14)
            run.bold = True

        # Address grid: For Account of Consignee / Deliver To Consignee
        address_table = doc.add_table(rows=1, cols=2)
        address_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        address_table.columns[0].width = Inches(3.75)
        address_table.columns[1].width = Inches(3.75)

        for i, label in enumerate(["For Account of Consignee", "Deliver To Consignee"]):
            cell = address_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            # Empty value field
            run = p.add_run("")
            run.font.size = Pt(12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Customer No. / Currency table
        customer_table = doc.add_table(rows=1, cols=2)
        customer_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        customer_table.columns[0].width = Inches(3.75)
        customer_table.columns[1].width = Inches(3.75)

        customer_labels = ["Customer No.", "Currency"]
        customer_values = ["", "EUR"]
        for i, (label, value) in enumerate(zip(customer_labels, customer_values)):
            cell = customer_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(18)
            run.bold = True

        # Payment Terms / Sales Area table
        payment_table = doc.add_table(rows=1, cols=2)
        payment_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        payment_table.columns[0].width = Inches(3.75)
        payment_table.columns[1].width = Inches(3.75)

        payment_labels = ["Payment Terms", "Sales Area"]
        payment_values = ["", ""]
        for i, (label, value) in enumerate(zip(payment_labels, payment_values)):
            cell = payment_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(18)
            run.bold = True

        # Incoterms / Bank Details table
        incoterms_table = doc.add_table(rows=1, cols=2)
        incoterms_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        incoterms_table.columns[0].width = Inches(3.75)
        incoterms_table.columns[1].width = Inches(3.75)

        incoterms_labels = ["Incoterms", "Bank Details"]
        incoterms_values = ["", ""]
        for i, (label, value) in enumerate(zip(incoterms_labels, incoterms_values)):
            cell = incoterms_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(18)
            run.bold = True

        # Shipping Details Grid (11 boxes in 2 rows)
        shipping_labels = [
            "Bill of Lading",
            "Loading Port / City",
            "Destination Port",
            "Despatch",
            "Mode / Vessel Name",
            "Container No.",
            "Voyage No.",
            "ETD",
            "ETA",
            "Shipping",
            "Seal No.",
        ]
        shipping_values = ["", "", "", "", "", "", "", "", "", "LCL", ""]

        # Create shipping details table with 6 columns (will wrap to next row)
        shipping_table = doc.add_table(rows=2, cols=6)
        shipping_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i in range(6):
            shipping_table.columns[i].width = Inches(1.25)

        for idx, (label, value) in enumerate(zip(shipping_labels, shipping_values)):
            row_idx = idx // 6
            col_idx = idx % 6
            if row_idx < 2 and col_idx < 6:
                cell = shipping_table.rows[row_idx].cells[col_idx]
                p = cell.add_paragraph()
                run = p.add_run(label)
                run.font.size = Pt(9)
                run.bold = True
                p = cell.add_paragraph()
                run = p.add_run(value)
                run.font.size = Pt(12)
                run.bold = True

        # Aggregate data from all shipments
        grand_total_cbm = 0.0
        grand_total_packages = 0
        grand_total_weight = 0.0
        grand_total_value = 0.0
        all_shipment_data = []
        all_shipment_types = []

        for shipment in shipments:
            pricing = calculate_invoice_totals(shipment)

            total_cbm = 0.0
            total_packages = 0
            total_weight = 0.0
            total_value = 0.0
            shipment_types = []

            shipment_items = []
            for idx, item in enumerate(pricing.get("parcel_calculations", [])):
                try:
                    cbm_value = float(item.get("cbm", 0) or 0)
                except (TypeError, ValueError):
                    cbm_value = 0.0
                total_cbm += cbm_value

                try:
                    weight_value = float(item.get("weight", 0) or 0)
                except (TypeError, ValueError):
                    weight_value = 0.0
                repeat_count = int(item.get("repeat_count", 1) or 1)
                total_weight += weight_value * repeat_count
                total_packages += repeat_count

                try:
                    price_value = float(item.get("price_by_weight", 0) or 0)
                except (TypeError, ValueError):
                    price_value = 0.0
                total_value += price_value

                # Get shipment type
                shipment_type = item.get("shipment_type")
                if (
                    not shipment_type
                    and shipment.parcels
                    and idx < len(shipment.parcels)
                ):
                    parcel = shipment.parcels[idx]
                    if isinstance(parcel, dict):
                        shipment_type = parcel.get("shipmentType") or parcel.get(
                            "shipment_type"
                        )
                if not shipment_type:
                    shipment_type = shipment.shipment_type

                if shipment_type:
                    shipment_types.append(shipment_type)

                shipment_items.append(
                    {
                        "shipment_number": shipment.shipment_number
                        or f"#{shipment.id}",
                        "product_name_en": item.get("product_name_en", ""),
                        "product_name_ar": item.get("product_name_ar", ""),
                        "hs_code": item.get("hs_code", ""),
                        "cbm": float(item.get("cbm", 0) or 0),
                        "repeat_count": int(item.get("repeat_count", 1) or 1),
                        "price_by_weight": float(item.get("price_by_weight", 0) or 0),
                        "price_per_kg": float(item.get("price_per_kg", 0) or 0),
                        "weight": float(item.get("weight", 0) or 0),
                        "shipment_type": shipment_type,
                    }
                )

            all_shipment_types.extend(shipment_types)

            grand_total_cbm += total_cbm
            grand_total_packages += total_packages
            grand_total_weight += total_weight
            grand_total_value += total_value

            unique_types = list(set(shipment_types))
            is_personal_only = len(unique_types) == 1 and unique_types[0] == "personal"
            is_commercial_only = (
                len(unique_types) == 1 and unique_types[0] == "commercial"
            )
            is_mixed = len(unique_types) > 1

            all_shipment_data.append(
                {
                    "shipment": shipment,
                    "pricing": pricing,
                    "items": shipment_items,
                    "total_cbm": total_cbm,
                    "total_packages": total_packages,
                    "total_weight": total_weight,
                    "total_value": total_value,
                    "is_personal_only": is_personal_only,
                    "is_commercial_only": is_commercial_only,
                    "is_mixed": is_mixed,
                }
            )

        # Product table - adjusted widths to fit within 7.5 inches
        product_table = doc.add_table(rows=1, cols=8)
        product_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        product_table.columns[0].width = Inches(0.75)  # Product Code
        product_table.columns[1].width = Inches(1.4)  # Client Name
        product_table.columns[2].width = Inches(1.8)  # Description
        product_table.columns[3].width = Inches(0.9)  # Shipment Type
        product_table.columns[4].width = Inches(0.75)  # Origin
        product_table.columns[5].width = Inches(0.75)  # HS Code
        product_table.columns[6].width = Inches(0.6)  # CBM
        product_table.columns[7].width = Inches(0.55)  # Unit

        headers = [
            "Product Code",
            "Client Name",
            "Description",
            "Shipment Type",
            "Origin",
            "HS Code",
            "CBM",
            "Unit",
        ]
        for i, header in enumerate(headers):
            cell = product_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell, "1F2937")
            run.font.color.rgb = RGBColor(255, 255, 255)

        for shipment_info in all_shipment_data:
            for item in shipment_info["items"]:
                row = product_table.add_row()
                row.cells[0].text = item["shipment_number"]
                row.cells[1].text = shipment_info["shipment"].sender_name or ""

                desc_text = item["product_name_en"]
                if item["product_name_ar"]:
                    desc_text += f"\n{item['product_name_ar']}"
                row.cells[2].text = desc_text

                shipment_type = item["shipment_type"]
                if shipment_type == "personal":
                    row.cells[3].text = "Personal / شخصي"
                elif shipment_type == "commercial":
                    row.cells[3].text = "Commercial / تجاري"
                else:
                    row.cells[3].text = "-"

                row.cells[4].text = ""
                row.cells[5].text = item["hs_code"] or ""
                row.cells[6].text = f"{item['cbm']:.3f}"
                row.cells[7].text = "PCS"

                for i in [3, 4, 5, 6, 7]:
                    row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Second product table - adjusted widths to fit within 7.5 inches
        product_table2 = doc.add_table(rows=1, cols=5)
        product_table2.alignment = WD_TABLE_ALIGNMENT.LEFT
        product_table2.columns[0].width = Inches(1.3)  # Qty
        product_table2.columns[1].width = Inches(1.5)  # Unit Price
        product_table2.columns[2].width = Inches(1.5)  # Value
        product_table2.columns[3].width = Inches(1.5)  # Gross Wt
        product_table2.columns[4].width = Inches(1.7)  # Nett Wt

        headers2 = ["Qty", "Unit Price", "Value", "Gross Wt", "Nett Wt"]
        for i, header in enumerate(headers2):
            cell = product_table2.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell, "1F2937")
            run.font.color.rgb = RGBColor(255, 255, 255)

        for shipment_info in all_shipment_data:
            for item in shipment_info["items"]:
                row = product_table2.add_row()
                row.cells[0].text = str(item["repeat_count"])

                shipment_type = item["shipment_type"]
                if shipment_type == "commercial":
                    row.cells[1].text = ""
                    row.cells[2].text = ""
                elif shipment_type == "personal":
                    row.cells[1].text = "€0.00"
                    row.cells[2].text = "€0.00"
                else:
                    if item["price_per_kg"]:
                        row.cells[1].text = f"€{item['price_per_kg']:.2f}"
                    row.cells[2].text = f"€{item['price_by_weight']:.2f}"

                row.cells[3].text = f"{item['weight']:.2f}"
                row.cells[4].text = ""

                for i in [1, 2, 3, 4]:
                    row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Totals table - adjusted widths to fit within 7.5 inches
        totals_table = doc.add_table(rows=1, cols=3)
        totals_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        totals_table.columns[0].width = Inches(2.5)
        totals_table.columns[1].width = Inches(2.5)
        totals_table.columns[2].width = Inches(2.5)

        totals_labels = ["SALES TAX", "PACKAGES", "TOTAL INVOICE VALUE"]
        totals_values = ["", str(grand_total_packages), f"EUR {grand_total_value:.2f}"]

        for i, (label, value) in enumerate(zip(totals_labels, totals_values)):
            cell = totals_table.rows[0].cells[i]
            p = cell.add_paragraph()
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.bold = True
            p = cell.add_paragraph()
            run = p.add_run(value)
            run.font.size = Pt(18)
            run.bold = True

        # Footer notes
        p = doc.add_paragraph()
        run = p.add_run("A - Legal & Customs Declaration\n")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(
            "MEDO-B2B EU acts solely as a freight consolidator and export agent on behalf of multiple clients. "
            "The declared values are for customs declaration purposes only and do not represent sales or transfer of ownership. "
            "This shipment is handled under EU Export Compliance (EX-A). MEDO-B2B EU bears no ownership or sales relation "
            "to the goods. All clients have submitted their own declaration forms confirming that their items are personal "
            "or commercial as described.\n\n"
        )

        run = p.add_run("B - Freight & Service Terms\n")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(
            "Payment is due only for freight and handling services, by bank transfer or cash, 100% prepaid before departure. "
            "The consignee acknowledges that all goods were inspected and accepted before loading. No return, refund, or "
            "complaint is accepted after the container has been sealed and departed. MEDO-B2B EU provides transport "
            "services under CIF Incoterms (Cost, Insurance, Freight) unless otherwise agreed in writing. Our general "
            "terms and conditions apply to all shipments and can be provided upon request."
        )

        # Signature section - adjusted widths to fit within 7.5 inches
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        sig_table.columns[0].width = Inches(3.75)
        sig_table.columns[1].width = Inches(3.75)

        for i in range(2):
            cell = sig_table.rows[0].cells[i]
            p = cell.add_paragraph()
            if i == 0:
                run = p.add_run("Authorized Signature & Stamp")
            else:
                run = p.add_run("Date")
            run.bold = True
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for _ in range(5):
                cell.add_paragraph()

            p = cell.add_paragraph()
            p.add_run("_" * 30)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        logger.info(
            f"Successfully generated consolidated export invoice Word document for {len(shipments)} shipments"
        )
        return doc_bytes.getvalue()

    except Exception as e:
        logger.error(
            f"Error generating consolidated export invoice Word document: {str(e)}",
            exc_info=True,
        )
        raise


def generate_consolidated_packing_list_bulk_word(
    shipments: List[LCLShipment],
    language: str = "en",
    packing_list_number: Optional[str] = None,
) -> bytes:
    """
    Generate Consolidated Packing List Word document for multiple LCL shipments.

    Args:
        shipments: List of LCLShipment instances
        language: kept for future use (currently template is EN)
        packing_list_number: Optional packing list number

    Returns:
        Word document bytes (.docx)
    """
    # Reuse the consolidated packing list function
    return generate_consolidated_packing_list_word(shipments, language)


def generate_multiple_consolidated_packing_lists_word(
    shipments: List[LCLShipment],
    language: str = "en",
    max_cbm: float = 65.0,
    max_weight_kg: float = 24000.0,
) -> bytes:
    """
    Generate multiple consolidated packing lists split by CBM and weight limits,
    and return them as a ZIP file containing Word documents.

    Args:
        shipments: List of LCLShipment instances
        language: Language for packing lists (default: 'en')
        max_cbm: Maximum CBM per packing list (default: 65.0)
        max_weight_kg: Maximum weight per packing list (default: 24000.0)

    Returns:
        ZIP file bytes containing all Word documents
    """
    import zipfile
    from datetime import datetime

    try:
        groups = split_shipments_by_limits(shipments, max_cbm, max_weight_kg)

        if not groups:
            raise ValueError("No shipment groups created")

        zip_buffer = io.BytesIO()
        date_str = datetime.now().strftime("%Y%m%d")

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for idx, group in enumerate(groups, start=1):
                packing_list_number = f"PL-{date_str}-{idx:03d}"
                docx_bytes = generate_consolidated_packing_list_bulk_word(
                    group,
                    language=language,
                    packing_list_number=packing_list_number,
                )
                filename = f"Consolidated-Packing-List-{packing_list_number}.docx"
                zip_file.writestr(filename, docx_bytes)

        zip_buffer.seek(0)
        zip_bytes = zip_buffer.getvalue()

        logger.info(
            f"Successfully generated {len(groups)} packing lists and compressed into ZIP file"
        )
        return zip_bytes

    except Exception as e:
        logger.error(
            f"Error generating multiple consolidated packing lists: {str(e)}",
            exc_info=True,
        )
        raise


def generate_multiple_consolidated_invoices_word(
    shipments: List[LCLShipment],
    language: str = "en",
    max_cbm: float = 65.0,
    max_weight_kg: float = 24000.0,
) -> bytes:
    """
    Generate multiple consolidated export invoices split by CBM and weight limits,
    and return them as a ZIP file containing Word documents.

    Args:
        shipments: List of LCLShipment instances
        language: Language for invoices (default: 'en')
        max_cbm: Maximum CBM per invoice (default: 65.0)
        max_weight_kg: Maximum weight per invoice (default: 24000.0)

    Returns:
        ZIP file bytes containing all Word documents
    """
    import zipfile
    from datetime import datetime

    try:
        groups = split_shipments_by_limits(shipments, max_cbm, max_weight_kg)

        if not groups:
            raise ValueError("No shipment groups created")

        zip_buffer = io.BytesIO()
        date_str = datetime.now().strftime("%Y%m%d")

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for idx, group in enumerate(groups, start=1):
                invoice_number = f"INV-{date_str}-{idx:03d}"
                docx_bytes = generate_consolidated_export_invoice_bulk_word(
                    group, language=language, invoice_number=invoice_number
                )
                filename = f"Consolidated-Export-Invoice-{invoice_number}.docx"
                zip_file.writestr(filename, docx_bytes)

        zip_buffer.seek(0)
        zip_bytes = zip_buffer.getvalue()

        logger.info(
            f"Successfully generated {len(groups)} invoices and compressed into ZIP file"
        )
        return zip_bytes

    except Exception as e:
        logger.error(
            f"Error generating multiple consolidated invoices: {str(e)}", exc_info=True
        )
        raise
